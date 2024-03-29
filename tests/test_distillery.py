import datetime
import glob
import inspect
import os
import random
import shutil
import string
import subprocess
import sys
import tempfile
import time
import urllib.request

import boto3
import git
import pytest

from decouple import config
from docx import Document
from essential_generators import DocumentGenerator
from gtts import gTTS
from PIL import Image, ImageDraw, ImageFont
from playwright.sync_api import Page, expect


@pytest.fixture(scope="session")
def browser_context_args(browser_context_args):
    return {
        **browser_context_args,
        "http_credentials": {
            "username": config("DISTILLERY_BASIC_AUTH_USERNAME", default=""),
            "password": config("DISTILLERY_BASIC_AUTH_PASSWORD", default=""),
        },
        # "record_video_dir": "tests/_output/",
    }


@pytest.fixture(autouse=True)
def set_timeout(page: Page):
    page.set_default_timeout(60000)
    return page


@pytest.fixture(autouse=True)
def distillery_0000_reset_files():
    for d in glob.glob(os.path.join(config("INITIAL_ORIGINAL_FILES"), "*/")):
        if d.split("/")[-1].split("-")[0] == "item":
            shutil.move(d, config("TEST_FILES", default="tests/files"))
    for batch in glob.glob(os.path.join(config("BATCH_SETS_DIRECTORY"), "*/")):
        for stage in glob.glob(os.path.join(batch, "*/")):
            for i in glob.glob(os.path.join(stage, "*")):
                if i.split("/")[-1].split("-")[0] == "item":
                    shutil.move(i, config("TEST_FILES", default="tests/files"))
        os.system(f"/bin/rm -r {batch}")
    for d in glob.glob(os.path.join(config("WORK_PRESERVATION_FILES"), "*/")):
        os.system(f"/bin/rm -r {d}")
    return


@pytest.fixture
def asnake_client():
    from asnake.client import ASnakeClient

    asnake_client = ASnakeClient(
        baseurl=config("ASPACE_API_URL"),
        username=config("ASPACE_USERNAME"),
        password=config("ASPACE_PASSWORD"),
    )
    asnake_client.authorize()
    return asnake_client


@pytest.fixture
def s3_client():
    s3_client = boto3.client(
        "s3",
        region_name=config("DISTILLERY_AWS_REGION", default="us-west-2"),
        aws_access_key_id=config("DISTILLERY_AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=config("DISTILLERY_AWS_SECRET_ACCESS_KEY"),
    )
    return s3_client


@pytest.fixture(scope="session")
def timestamp():
    return str(time.time())


# NOTE cleaning up previous tests and then setting up for new tests is not the
# recommended way to do Arrange and Cleanup in pytest, but it allows manually
# viewing the outcome of tests in ArchivesSpace


def delete_archivesspace_test_records(asnake_client, resource_identifer):
    """Delete any existing test records."""

    def delete_related_records(uri):
        result = asnake_client.get(uri).json()
        for instance in result["instances"]:
            if instance.get("digital_object"):
                print(f'🐞 DELETING DIGITAL_OBJECT {instance["digital_object"]["ref"]}')
                digital_object_delete_response = asnake_client.delete(
                    instance["digital_object"]["ref"]
                )
            if instance.get("sub_container") and instance["sub_container"].get(
                "top_container"
            ):
                print(
                    f'🐞 DELETING TOP_CONTAINER {instance["sub_container"]["top_container"]["ref"]}'
                )
                top_container_delete_response = asnake_client.delete(
                    instance["sub_container"]["top_container"]["ref"]
                )
        for linked_agent in result["linked_agents"]:
            print(f'🐞 DELETING LINKED_AGENT {linked_agent["ref"]}')
            linked_agent_delete_response = asnake_client.delete(linked_agent["ref"])

    def recursive_delete(node, resource_ref):
        if node["child_count"] > 0:
            children = asnake_client.get(
                f'{resource_ref}/tree/waypoint?offset=0&parent_node={node["uri"]}'
            ).json()
            for child in children:
                delete_related_records(child["uri"])
                recursive_delete(child, resource_ref)

    resource_find_by_id_results = asnake_client.get(
        "/repositories/2/find_by_id/resources",
        params={"identifier[]": [f'["{resource_identifer}"]']},
    ).json()
    print(f'🐞 RESOURCE_FIND_BY_ID_RESULTS["RESOURCES"]: {resource_find_by_id_results}')
    for result in resource_find_by_id_results["resources"]:
        delete_related_records(result["ref"])
        resource_tree = asnake_client.get(f'{result["ref"]}/tree/root').json()
        if resource_tree["waypoints"] > 1:
            raise Exception("Test resource has more than one waypoint.")
        resource_children = resource_tree["precomputed_waypoints"][""]["0"]
        for child in resource_children:
            delete_related_records(child["uri"])
            recursive_delete(child, result["ref"])
        print(f'🐞 DELETING RESOURCE {result["ref"]}')
        asnake_client.delete(result["ref"])


def move_test_files_to_initial_original_files_directory(test_name):
    shutil.move(
        os.path.join(config("TEST_FILES", default="tests/files"), f"item-{test_name}"),
        config("INITIAL_ORIGINAL_FILES"),
    )


def invalidate_cloudfront_path(path="/*", caller_reference=None):
    if not caller_reference:
        caller_reference = str(time.time())
    print(f"🐞 CLOUDFRONT INVALIDATION PATH: {path}")
    print(f"🐞 CLOUDFRONT INVALIDATION CALLER_REFERENCE: {caller_reference}")
    cloudfront_client = boto3.client(
        "cloudfront",
        aws_access_key_id=config("DISTILLERY_AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=config("DISTILLERY_AWS_SECRET_ACCESS_KEY"),
    )
    response = cloudfront_client.create_invalidation(
        DistributionId=config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID"),
        InvalidationBatch={
            "Paths": {"Quantity": 1, "Items": [path]},
            "CallerReference": caller_reference,
        },
    )
    print(f"🐞 CLOUDFRONT INVALIDATION RESPONSE: {str(response)}")
    waiter = cloudfront_client.get_waiter("invalidation_completed")
    print("🐞 WAITING ON CLOUDFRONT INVALIDATION")
    waiter.wait(
        DistributionId=config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID"),
        Id=response["Invalidation"]["Id"],
    )
    print("🐞 CLOUDFRONT INVALIDATION COMPLETE")


def create_archivesspace_test_resource(asnake_client, test_name, test_id):
    resource = {}
    # required
    resource["title"] = f'{test_name.capitalize().replace("_", " ")}'
    # NOTE `id_0` is limited to 50 characters
    resource["id_0"] = test_id
    resource["level"] = "collection"
    resource["finding_aid_language"] = "eng"
    resource["finding_aid_script"] = "Latn"
    resource["lang_materials"] = [
        {"language_and_script": {"language": "eng", "script": "Latn"}}
    ]
    resource["dates"] = [
        {
            "label": "creation",
            "date_type": "single",
            "begin": str(datetime.date.today()),
        }
    ]
    resource["extents"] = [{"portion": "whole", "number": "1", "extent_type": "boxes"}]
    # optional
    resource["publish"] = True
    # post
    resource_post_response = asnake_client.post(
        "/repositories/2/resources", json=resource
    )
    return resource_post_response


def create_archivesspace_test_archival_object_item(
    asnake_client, test_name, test_id, resource_uri, **kwargs
):
    item = {}
    # required
    item["title"] = f"Item {test_id}"
    item["level"] = "item"
    item["resource"] = {"ref": resource_uri}
    # optional
    item["component_id"] = f"item-{test_name}"
    item["publish"] = True
    # customizations
    if kwargs.get("customizations"):
        for key, value in kwargs["customizations"].items():
            if key == "linked_agents":
                item[key] = []
                for agent in value:
                    ref = create_archivesspace_test_agent_person(
                        asnake_client, test_id, agent
                    )
                    item[key].append(
                        {
                            "ref": ref.json()["uri"],
                            "role": agent.get("role", "creator"),
                            "relator": agent.get("relator"),
                        }
                    )
            else:
                item[key] = value
    # post
    item_post_response = asnake_client.post(
        "/repositories/2/archival_objects", json=item
    )
    return item_post_response, item["component_id"]


def create_archivesspace_test_archival_object_series(
    asnake_client, test_id, resource_uri, **kwargs
):
    series = {}
    # required
    series["title"] = f"[Series] {test_id}"
    series["level"] = "series"
    series["resource"] = {"ref": resource_uri}
    # optional
    series["component_id"] = f"series-{test_id}"
    series["publish"] = True
    # customizations
    if kwargs.get("customizations"):
        for key, value in kwargs["customizations"].items():
            series[key] = value
    # post
    series_post_response = asnake_client.post(
        "/repositories/2/archival_objects", json=series
    )
    return series_post_response


def create_archivesspace_test_archival_object_subseries(
    asnake_client, test_id, resource_uri, **kwargs
):
    subseries = {}
    # required
    subseries["title"] = f"[Sub-Series] {test_id}"
    subseries["level"] = "subseries"
    subseries["resource"] = {"ref": resource_uri}
    # optional
    subseries["component_id"] = f"subseries-{test_id}"
    subseries["publish"] = True
    # customizations
    if kwargs.get("customizations"):
        for key, value in kwargs["customizations"].items():
            subseries[key] = value
    # post
    subseries_post_response = asnake_client.post(
        "/repositories/2/archival_objects", json=subseries
    )
    return subseries_post_response


def create_archivesspace_test_agent_person(asnake_client, test_id, customizations):
    rest_of_name = "{} {}".format(
        "Published" if customizations.get("publish") else "Unpublished",
        customizations.get("role").capitalize(),
    )
    person = {
        "names": [
            {
                "name_order": "inverted",
                "primary_name": test_id.capitalize(),
                "rest_of_name": rest_of_name,
                "sort_name": f"{test_id.capitalize()}, {rest_of_name}",
            }
        ],
        "publish": customizations.get("publish", False),
    }
    agent_person_create_response = asnake_client.post("/agents/people", json=person)
    print(
        "🐞 agent_person_create_response",
        agent_person_create_response.json(),
    )
    return agent_person_create_response


def create_archivesspace_test_digital_object(asnake_client, test_name, test_id):
    digital_object = {}
    # required
    digital_object["title"] = f"Item {test_id}"
    digital_object["digital_object_id"] = f"item-{test_name}"
    # optional
    digital_object["file_versions"] = [
        {"file_uri": f"http://example.org/item-{test_id}", "publish": True},
        {
            "file_uri": f"http://example.org/item-{test_id}/thumbnail",
            "publish": True,
            "xlink_show_attribute": "embed",
        },
    ]
    digital_object["publish"] = True
    # post
    digital_object_post_response = asnake_client.post(
        "/repositories/2/digital_objects", json=digital_object
    )
    return digital_object_post_response, digital_object["digital_object_id"]


def unblock_archivesspace_ports():
    subprocess.run(
        [
            config("WORK_SSH_CMD", default="/usr/bin/ssh"),
            "-i",
            config("ARCHIVESSPACE_SSH_KEY"),
            "-p",
            config("ARCHIVESSPACE_SSH_PORT", default="22"),
            f'{config("ARCHIVESSPACE_SSH_USER")}@{config("ARCHIVESSPACE_SSH_HOST")}',
            *config(
                "ARCHIVESSPACE_SIMULATE_ONLINE_CMD",
                default="sudo ufw disable",
            ).split(),
        ],
        check=True,
        capture_output=True,
    )


def publish_to_web(page: Page, **kwargs):
    page.goto(config("DISTILLERY_BASE_URL"))
    page.locator(f'input[value="access"]').click()
    page.locator(f'input[value="{kwargs.get("file_versions_op", "fail")}"]').check()
    page.locator(f'input[value="{kwargs.get("thumbnail_label", "sequence")}"]').check()
    page.get_by_role("button", name="Validate").click()
    page.locator("summary").click()
    if kwargs.get("outcome", "success") == "failure":
        expect(page.locator("p")).to_have_text(
            "❌ Something went wrong. View the details for more information.",
            timeout=kwargs.get("timeout", 60000),
        )
        return page
    success = page.locator("p").get_by_text("✅ Successfully")
    validated = page.locator("p").get_by_text(
        "✅ Successfully validated metadata, files, and destinations."
    )
    processed = page.locator("p").get_by_text(
        "✅ Successfully processed metadata and files."
    )
    error = page.locator("p").get_by_text(
        "❌ Something went wrong. View the details for more information."
    )
    expect(success.or_(error)).to_be_visible(timeout=kwargs.get("timeout", 60000))
    expect(validated).to_be_visible()
    page.get_by_role("button", name="Run").click()
    page.locator("summary").click()
    expect(success.or_(error)).to_be_visible(timeout=kwargs.get("timeout", 60000))
    expect(processed).to_be_visible()


def run_distillery(
    page: Page,
    destinations,
    file_versions_op="fail",
    thumbnail_label="sequence",
    outcome="success",
    timeout=60000,
):
    page.goto(config("DISTILLERY_BASE_URL"))
    for destination in destinations:
        page.locator(f'input[value="{destination}"]').click()
        if destination == "access":
            page.locator(f'input[value="{file_versions_op}"]').check()
            page.locator(f'input[value="{thumbnail_label}"]').check()
    page.get_by_role("button", name="Validate").click()
    page.locator("summary").click()
    if outcome == "failure":
        expect(page.locator("p")).to_have_text(
            "❌ Something went wrong. View the details for more information.",
            timeout=timeout,
        )
        return page
    success = page.locator("p").get_by_text("✅ Successfully")
    validated = page.locator("p").get_by_text(
        "✅ Successfully validated metadata, files, and destinations."
    )
    processed = page.locator("p").get_by_text(
        "✅ Successfully processed metadata and files."
    )
    error = page.locator("p").get_by_text(
        "❌ Something went wrong. View the details for more information."
    )
    expect(success.or_(error)).to_be_visible(timeout=timeout)
    expect(validated).to_be_visible()
    page.get_by_role("button", name="Run").click()
    page.locator("summary").click()
    expect(success.or_(error)).to_be_visible(timeout=timeout)
    expect(processed).to_be_visible()


def run_alchemist_regenerate(
    page: Page,
    content_attributes,
    regenerate,
    simulate_archivesspace_offline,
    timeout=300000,
):
    page.goto("/".join([config("DISTILLERY_BASE_URL").rstrip("/"), "alchemist"]))
    if regenerate == "one":
        page.get_by_label("Regenerate files for one item").check()
        assert content_attributes[0]["component_id"]
        page.get_by_label("Component Unique Identifier").fill(
            content_attributes[0]["component_id"]
        )
    if regenerate == "collection":
        page.get_by_label("Regenerate files for a collection").check()
        assert content_attributes[0]["id_0"]
        page.get_by_label("Collection Identifier").fill(content_attributes[0]["id_0"])
    elif regenerate == "all":
        page.get_by_label("Regenerate files for all items").check()
    page.get_by_role("button", name="Regenerate").click()
    page.locator("summary").click()
    if simulate_archivesspace_offline:
        print("🐞 SIMULATING ARCHIVESSPACE OFFLINE")
        subprocess.run(
            [
                config("WORK_SSH_CMD", default="/usr/bin/ssh"),
                "-i",
                config("ARCHIVESSPACE_SSH_KEY"),
                "-p",
                config("ARCHIVESSPACE_SSH_PORT", default="22"),
                f'{config("ARCHIVESSPACE_SSH_USER")}@{config("ARCHIVESSPACE_SSH_HOST")}',
                *config(
                    "ARCHIVESSPACE_SIMULATE_OFFLINE_CMD",
                    default="sudo ufw allow 22 && sudo ufw --force enable",
                ).split(),
            ],
            check=True,
            capture_output=True,
        )
        # NOTE we need to wait for the CloudFront invalidations to complete and
        # then wait for requests to time out when looking for ArchivesSpace
        print("🐞 SLEEPING 5 MINUTES...")
        time.sleep(60)
        expect(page.get_by_text("❌ Something went wrong.")).not_to_be_visible()
        print("🐞 ... 4 MINUTES")
        time.sleep(60)
        expect(page.get_by_text("❌ Something went wrong.")).not_to_be_visible()
        print("🐞 ... 3 MINUTES")
        time.sleep(60)
        expect(page.get_by_text("❌ Something went wrong.")).not_to_be_visible()
        print("🐞 ... 2 MINUTES")
        time.sleep(60)
        expect(page.get_by_text("❌ Something went wrong.")).not_to_be_visible()
        print("🐞 ... 1 MINUTE")
        time.sleep(60)
        expect(page.get_by_text("❌ Something went wrong.")).not_to_be_visible()
        print("🐞 SIMULATING ARCHIVESSPACE ONLINE")
        unblock_archivesspace_ports()
    expect(page.get_by_text("❌ Something went wrong.")).not_to_be_visible()
    expect(page.locator("p")).to_contain_text(f"✅ Regenerated", timeout=300000)


def run_oralhistories_add(page: Page, file, outcome="success"):
    page.goto("/".join([config("DISTILLERY_BASE_URL").rstrip("/"), "oralhistories"]))
    page.locator("#file").set_input_files(file)
    page.get_by_role("button", name="Upload").click()
    if outcome == "failure":
        expect(page.frame_locator("iframe").locator("body")).to_contain_text(
            "❌", timeout=60000
        )
        return page
    expect(page.frame_locator("iframe").locator("body")).to_contain_text(
        "🏁", timeout=60000
    )


def run_oralhistories_publish(page: Page, item_component_id):
    page.goto("/".join([config("DISTILLERY_BASE_URL").rstrip("/"), "oralhistories"]))
    page.locator("#component_id_publish").fill(item_component_id)
    page.get_by_role("button", name="Publish Changes").click()
    expect(page.frame_locator("iframe").locator("body")).to_contain_text(
        "🏁", timeout=60000
    )


def run_oralhistories_update(page: Page, item_component_id):
    page.goto("/".join([config("DISTILLERY_BASE_URL").rstrip("/"), "oralhistories"]))
    page.locator("#component_id_update").fill(item_component_id)
    page.get_by_role("button", name="Update Metadata").click()
    expect(page.frame_locator("iframe").locator("body")).to_contain_text(
        "🏁", timeout=60000
    )


def format_alchemist_item_uri(test_name, test_id):
    return "/".join(
        [
            config("ALCHEMIST_BASE_URL").rstrip("/"),
            config("ALCHEMIST_URL_PREFIX"),
            test_id,
            f"item-{test_name}",
        ]
    )


def wait_for_oralhistories_generated_files(git_repo, attempts=3, sleep_time=30):
    attempts = attempts
    while True:
        git_repo.remotes.origin.pull()
        print(f"🐞 git_repo.head.commit.message: {git_repo.head.commit.message.strip()}")
        if "generated files" in git_repo.head.commit.message:
            return True
        elif attempts > 0:
            print(f"🐞 waiting {sleep_time} seconds")
            time.sleep(sleep_time)
        else:
            return False
        attempts -= 1


def generate_image_file(file_stem, **kwargs):
    tmp_file, headers = urllib.request.urlretrieve(
        config("TEST_IMG_URL", default="https://picsum.photos/1600/1200")
    )
    img = Image.open(tmp_file)
    drw = ImageDraw.Draw(img)
    fnt = ImageFont.truetype(
        config(
            "TEST_IMG_FONT",
            default="/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ),
        144,
    )
    drw.text(
        (800, 600),
        file_stem.split("_")[-1],
        fill="white",
        font=fnt,
        anchor="mm",
        stroke_width=10,
        stroke_fill="black",
    )
    if "directory" in kwargs:
        file = os.path.join(
            config("INITIAL_ORIGINAL_FILES"),
            f'item-{kwargs["directory"]}',
            f"{file_stem}.tif",
        )
        os.makedirs(os.path.dirname(file), exist_ok=True)
    else:
        file = os.path.join(config("INITIAL_ORIGINAL_FILES"), f"item-{file_stem}.tif")
    img.save(file)
    img.close()


def generate_video_file(test_name, **kwargs):
    if "directory" in kwargs:
        output_file = os.path.join(
            config("INITIAL_ORIGINAL_FILES"),
            f'item-{kwargs["directory"]}',
            f"item-{test_name}.mp4",
        )
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
    else:
        output_file = os.path.join(
            config("INITIAL_ORIGINAL_FILES"), f"item-{test_name}.mp4"
        )
    timecode_text = r"text='timestamp:%{pts\:hms}': x=(w-text_w)/2: y=100: font=Mono: fontsize=64: fontcolor=Yellow: box=1: boxcolor=Black: boxborderw=10"
    testname_text = f'text={test_name.split("_")[-1]}: x=(w-text_w)/2: y=(h-text_h)/2: fontsize=196: fontcolor=White: shadowx=3: shadowy=3'
    subprocess.run(
        [
            config("WORK_FFMPEG_CMD"),
            "-f",
            "lavfi",
            "-i",
            "smptebars=s=1280x720",
            "-vf",
            f"drawtext={timecode_text}, drawtext={testname_text}",
            "-c:v",
            "libx264",
            "-r",
            "30",
            "-frames:v",
            "90",
            "-y",
            output_file,
        ],
        check=True,
    )


def generate_audio_file(test_name, **kwargs):
    audio = gTTS(text=test_name.split("_")[-1])
    if "directory" in kwargs:
        file = os.path.join(
            config("INITIAL_ORIGINAL_FILES"),
            f'item-{kwargs["directory"]}',
            f"item-{test_name}.mp3",
        )
        os.makedirs(os.path.dirname(file), exist_ok=True)
    else:
        file = os.path.join(config("INITIAL_ORIGINAL_FILES"), f"item-{test_name}.mp3")
    audio.save(file)


def generate_filesystem_files(test_name, **kwargs):
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    def _generate_audio_file(file_stem, **kwargs):
        audio = gTTS(text=file_stem.split("_")[-1])
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.mp3",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.mp3"
            )
        audio.save(output_file)
        print(f"🐞 GENERATED FILE {output_file}")

    def _generate_video_file(file_stem, **kwargs):
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.mp4",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.mp4"
            )
        timecode_text = r"text='timestamp:%{pts\:hms}': x=(w-text_w)/2: y=100: font=Mono: fontsize=64: fontcolor=Yellow: box=1: boxcolor=Black: boxborderw=10"
        testname_text = f'text={file_stem.split("_")[-1]}: x=(w-text_w)/2: y=(h-text_h)/2: fontsize=196: fontcolor=White: shadowx=3: shadowy=3'
        subprocess.run(
            [
                config("WORK_FFMPEG_CMD"),
                "-f",
                "lavfi",
                "-i",
                "smptebars=s=1280x720",
                "-vf",
                f"drawtext={timecode_text}, drawtext={testname_text}",
                "-c:v",
                "libx264",
                "-r",
                "30",
                "-frames:v",
                "90",
                "-y",
                output_file,
            ],
            check=True,
        )
        print(f"🐞 GENERATED FILE {output_file}")

    def _generate_transcript_file(file_stem, **kwargs):
        generate = DocumentGenerator()
        document = Document()
        document.add_heading(file_stem.replace("_", " "), 0)
        document.add_page_break()
        document.add_paragraph(generate.paragraph())
        document.add_heading(generate.sentence(), level=1)
        document.add_heading(generate.sentence(), level=2)
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        document.add_heading(generate.sentence(), level=2)
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_heading(generate.sentence(), level=1)
        document.add_heading(generate.sentence(), level=2)
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        document.add_heading(generate.sentence(), level=2)
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        document.add_paragraph(generate.paragraph())
        paragraph = document.add_paragraph()
        paragraph.add_run("Question: ").bold = True
        paragraph.add_run(generate.sentence())
        paragraph = document.add_paragraph()
        paragraph.add_run("Answer: ").bold = True
        paragraph.add_run(generate.paragraph())
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.docx",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.docx"
            )
        document.save(output_file)
        print(f"🐞 GENERATED FILE {output_file}")
        return output_file

    def _generate_image_file(file_stem, **kwargs):
        tmp_file, headers = urllib.request.urlretrieve(
            config("TEST_IMG_URL", default="https://picsum.photos/1600/1200")
        )
        img = Image.open(tmp_file)
        drw = ImageDraw.Draw(img)
        fnt = ImageFont.truetype(
            config(
                "TEST_IMG_FONT",
                default="/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            ),
            144,
        )
        drw.text(
            (800, 600),
            file_stem.split("_")[-1],
            fill="white",
            font=fnt,
            anchor="mm",
            stroke_width=10,
            stroke_fill="black",
        )
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.tif",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.tif"
            )
        img.save(output_file)
        img.close()
        print(f"🐞 GENERATED FILE {output_file}")

    def _generate_digital_files(component_id, **kwargs):
        output_files = []
        if "mixed" in component_id:
            # NOTE "mixed" will always generate one each of audio, video, image
            # regardless of archival_object_count value
            if component_id.endswith("1"):
                _generate_audio_file(component_id, **kwargs)
            elif component_id.endswith("2"):
                _generate_video_file(component_id, **kwargs)
            elif component_id.endswith("3"):
                _generate_image_file(component_id, **kwargs)
        elif "audio" in component_id:
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    _generate_audio_file(
                        "{}{}".format(
                            "".join(
                                random.choices(
                                    string.ascii_lowercase + string.digits, k=6
                                )
                            ),
                            str(i).zfill(2),
                        ),
                        directory=component_id,
                        **kwargs,
                    )
            else:
                _generate_audio_file(component_id, **kwargs)
        elif "video" in component_id:
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    _generate_video_file(
                        "{}{}".format(
                            "".join(
                                random.choices(
                                    string.ascii_lowercase + string.digits, k=6
                                )
                            ),
                            str(i).zfill(2),
                        ),
                        directory=component_id,
                        **kwargs,
                    )
            else:
                _generate_video_file(component_id, **kwargs)
        elif "transcript" in component_id:
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    output_files.extend(
                        _generate_transcript_file(
                            "{}{}".format(
                                "".join(
                                    random.choices(
                                        string.ascii_lowercase + string.digits, k=6
                                    )
                                ),
                                str(i).zfill(2),
                            ),
                            directory=component_id,
                            **kwargs,
                        )
                    )
            else:
                output_files.append(_generate_transcript_file(component_id, **kwargs))
        else:  # create images by default
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    _generate_image_file(
                        "{}{}".format(
                            "".join(
                                random.choices(
                                    string.ascii_lowercase + string.digits, k=6
                                )
                            ),
                            str(i).zfill(2),
                        ),
                        directory=component_id,
                        **kwargs,
                    )
            else:
                _generate_image_file(component_id, **kwargs)
        return output_files

    output_files = []
    for j in range(1, kwargs["collection_count"] + 1):
        for i in range(1, kwargs["archival_object_count"] + 1):
            output_files.extend(
                _generate_digital_files(
                    f'{kwargs["level"]}__{test_name}xx{j}xx{i}', **kwargs
                )
            )
    return output_files


def delete_s3_preservation_objects(s3_client, test_id):
    s3_response = s3_client.list_objects_v2(
        Bucket=config("PRESERVATION_BUCKET"), Prefix=test_id
    )
    print("🐞 s3_client.list_objects_v2", s3_response)
    if s3_response.get("Contents"):
        s3_keys = [{"Key": s3_object["Key"]} for s3_object in s3_response["Contents"]]
        s3_response = s3_client.delete_objects(
            Bucket=config("PRESERVATION_BUCKET"), Delete={"Objects": s3_keys}
        )


def generate_archivesspace_records(test_name, asnake_client, **kwargs):
    # NOTE "mixed" will always generate one each of audio, video, image
    # regardless of archival_object_count value
    if "mixed" in test_name:
        kwargs["archival_object_count"] = 3

    def _create_archivesspace_resource(asnake_client, test_name, **kwargs):
        resource = {}
        # required
        # NOTE `id_0` is limited to 50 characters
        resource["title"] = f'{test_name.capitalize().replace("_", " ")} {kwargs["j"]}'
        resource["id_0"] = f'{test_name.split("_")[-1]}xx{kwargs["j"]}'
        resource["level"] = "collection"
        resource["finding_aid_language"] = "eng"
        resource["finding_aid_script"] = "Latn"
        resource["lang_materials"] = [
            {"language_and_script": {"language": "eng", "script": "Latn"}}
        ]
        resource["dates"] = [
            {
                "label": "creation",
                "date_type": "single",
                "begin": str(datetime.date.today()),
            }
        ]
        resource["extents"] = [
            {"portion": "whole", "number": "1", "extent_type": "boxes"}
        ]
        # optional
        resource["publish"] = True
        # post
        response = asnake_client.post("/repositories/2/resources", json=resource)
        print("🐞 RESOURCE", response.json())
        return {
            "uri": response.json()["uri"],
            "id_0": resource["id_0"],
        }

    def _create_archival_object(asnake_client, test_name, resource_uri, **kwargs):
        archival_object = {}
        archival_object["level"] = kwargs.get("level")
        archival_object["resource"] = {"ref": resource_uri}
        archival_object[
            "title"
        ] = f'{kwargs.get("level")} {test_name.split("_")[-1]} {kwargs["j"]} {kwargs["i"]}'
        archival_object[
            "component_id"
        ] = f'{kwargs.get("level")}__{test_name}xx{kwargs["j"]}xx{kwargs["i"]}'
        archival_object["publish"] = True
        # customizations
        if kwargs.get("customizations"):
            for key, value in kwargs["customizations"].items():
                if key == "linked_agents":
                    archival_object[key] = []
                    for agent in value:
                        ref = create_archivesspace_test_agent_person(
                            asnake_client, test_name.split("_")[-1], agent
                        )
                        archival_object[key].append(
                            {
                                "ref": ref.json()["uri"],
                                "role": agent.get("role", "creator"),
                                "relator": agent.get("relator"),
                            }
                        )
                else:
                    archival_object[key] = value
        # post
        response = asnake_client.post(
            "/repositories/2/archival_objects", json=archival_object
        )
        print(
            f'🐞 ARCHIVAL OBJECT {archival_object["component_id"]}',
            response.json(),
            f'{config("ASPACE_STAFF_URL").rstrip("/")}/resolve/readonly?uri={response.json()["uri"]}',
        )
        return {
            "uri": response.json()["uri"],
            "component_id": archival_object["component_id"],
        }

    def _create_ancestors(asnake_client, test_name, resource_uri, **kwargs):
        if "series" in kwargs.get("ancestors"):
            kwargs["level"] = "series"
            _series = _create_archival_object(
                asnake_client, test_name, resource_uri, **kwargs
            )
            if "subseries" in kwargs.get("ancestors"):
                kwargs["level"] = "subseries"
                _subseries = _create_archival_object(
                    asnake_client, test_name, resource_uri, **kwargs
                )
                asnake_client.post(
                    f'{_subseries["uri"]}/parent',
                    params={"parent": _series["uri"].rsplit("/")[-1], "position": 1},
                )
                if "file" in kwargs.get("ancestors"):
                    kwargs["level"] = "file"
                    _file = _create_archival_object(
                        asnake_client, test_name, resource_uri, **kwargs
                    )
                    asnake_client.post(
                        f'{_file["uri"]}/parent',
                        params={
                            "parent": _subseries["uri"].rsplit("/")[-1],
                            "position": 1,
                        },
                    )
                    return _file["uri"].rsplit("/")[-1]
                else:
                    return _subseries["uri"].rsplit("/")[-1]
            else:
                return _series["uri"].rsplit("/")[-1]
        elif "subseries" in kwargs.get("ancestors"):
            kwargs["level"] = "subseries"
            _subseries = _create_archival_object(
                asnake_client, test_name, resource_uri, **kwargs
            )
            if "file" in kwargs.get("ancestors"):
                kwargs["level"] = "file"
                _file = _create_archival_object(
                    asnake_client, test_name, resource_uri, **kwargs
                )
                asnake_client.post(
                    f'{_file["uri"]}/parent',
                    params={"parent": _subseries["uri"].rsplit("/")[-1], "position": 1},
                )
                return _file["uri"].rsplit("/")[-1]
            else:
                return _subseries["uri"].rsplit("/")[-1]
        elif "file" in kwargs.get("ancestors"):
            kwargs["level"] = "file"
            _file = _create_archival_object(
                asnake_client, test_name, resource_uri, **kwargs
            )
            return _file["uri"].rsplit("/")[-1]
        else:
            return None

    def _nest_archival_object(archival_object_uri, parent_id):
        asnake_client.post(
            f"{archival_object_uri}/parent", params={"parent": parent_id, "position": 1}
        )

    _content_attributes = []
    for j in range(1, kwargs.get("collection_count") + 1):
        _resource = _create_archivesspace_resource(asnake_client, test_name, j=j)
        parent_id = _create_ancestors(
            asnake_client,
            test_name,
            _resource["uri"],
            i=0,
            j=j,
            **kwargs,
        )
        for i in range(1, kwargs.get("archival_object_count") + 1):
            _archival_object = _create_archival_object(
                asnake_client,
                test_name,
                _resource["uri"],
                i=i,
                j=j,
                **kwargs,
            )
            if parent_id:
                _nest_archival_object(_archival_object["uri"], parent_id)
            _content_attributes.append(
                {
                    "archival_object_uri": _archival_object["uri"],
                    "component_id": _archival_object["component_id"],
                    "resource_uri": _resource["uri"],
                    "id_0": _resource["id_0"],
                }
            )
    return _content_attributes


def generate_content(test_name, asnake_client, **kwargs):
    # NOTE "mixed" will always generate one each of audio, video, image
    # regardless of archival_object_count value
    if "mixed" in test_name:
        kwargs["archival_object_count"] = 3

    def _create_archivesspace_resource(asnake_client, test_name, **kwargs):
        resource = {}
        # required
        # NOTE `id_0` is limited to 50 characters
        resource["title"] = f'{test_name.capitalize().replace("_", " ")} {kwargs["j"]}'
        resource["id_0"] = f'{test_name.split("_")[-1]}xx{kwargs["j"]}'
        resource["level"] = "collection"
        resource["finding_aid_language"] = "eng"
        resource["finding_aid_script"] = "Latn"
        resource["lang_materials"] = [
            {"language_and_script": {"language": "eng", "script": "Latn"}}
        ]
        resource["dates"] = [
            {
                "label": "creation",
                "date_type": "single",
                "begin": str(datetime.date.today()),
            }
        ]
        resource["extents"] = [
            {"portion": "whole", "number": "1", "extent_type": "boxes"}
        ]
        # optional
        resource["publish"] = True
        # post
        response = asnake_client.post("/repositories/2/resources", json=resource)
        print("🐞 RESOURCE", response.json())
        return {
            "uri": response.json()["uri"],
            "id_0": resource["id_0"],
        }

    def _create_archival_object(asnake_client, test_name, resource_uri, **kwargs):
        archival_object = {}
        archival_object["level"] = kwargs.get("level")
        archival_object["resource"] = {"ref": resource_uri}
        archival_object[
            "title"
        ] = f'{kwargs.get("level")} {test_name.split("_")[-1]} {kwargs["j"]} {kwargs["i"]}'
        archival_object[
            "component_id"
        ] = f'{kwargs.get("level")}__{test_name}xx{kwargs["j"]}xx{kwargs["i"]}'
        archival_object["publish"] = True
        # customizations
        if kwargs.get("customizations"):
            for key, value in kwargs["customizations"].items():
                if key == "linked_agents":
                    archival_object[key] = []
                    for agent in value:
                        ref = create_archivesspace_test_agent_person(
                            asnake_client, test_name.split("_")[-1], agent
                        )
                        archival_object[key].append(
                            {
                                "ref": ref.json()["uri"],
                                "role": agent.get("role", "creator"),
                                "relator": agent.get("relator"),
                            }
                        )
                else:
                    archival_object[key] = value
        # post
        response = asnake_client.post(
            "/repositories/2/archival_objects", json=archival_object
        )
        print(
            f'🐞 ARCHIVAL OBJECT {archival_object["component_id"]}',
            response.json(),
            f'{config("ASPACE_STAFF_URL").rstrip("/")}/resolve/readonly?uri={response.json()["uri"]}',
        )
        return {
            "uri": response.json()["uri"],
            "component_id": archival_object["component_id"],
        }

    def _create_ancestors(asnake_client, test_name, resource_uri, **kwargs):
        if "series" in kwargs.get("ancestors"):
            kwargs["level"] = "series"
            _series = _create_archival_object(
                asnake_client, test_name, resource_uri, **kwargs
            )
            if "subseries" in kwargs.get("ancestors"):
                kwargs["level"] = "subseries"
                _subseries = _create_archival_object(
                    asnake_client, test_name, resource_uri, **kwargs
                )
                asnake_client.post(
                    f'{_subseries["uri"]}/parent',
                    params={"parent": _series["uri"].rsplit("/")[-1], "position": 1},
                )
                if "file" in kwargs.get("ancestors"):
                    kwargs["level"] = "file"
                    _file = _create_archival_object(
                        asnake_client, test_name, resource_uri, **kwargs
                    )
                    asnake_client.post(
                        f'{_file["uri"]}/parent',
                        params={
                            "parent": _subseries["uri"].rsplit("/")[-1],
                            "position": 1,
                        },
                    )
                    return _file["uri"].rsplit("/")[-1]
                else:
                    return _subseries["uri"].rsplit("/")[-1]
            else:
                return _series["uri"].rsplit("/")[-1]
        elif "subseries" in kwargs.get("ancestors"):
            kwargs["level"] = "subseries"
            _subseries = _create_archival_object(
                asnake_client, test_name, resource_uri, **kwargs
            )
            if "file" in kwargs.get("ancestors"):
                kwargs["level"] = "file"
                _file = _create_archival_object(
                    asnake_client, test_name, resource_uri, **kwargs
                )
                asnake_client.post(
                    f'{_file["uri"]}/parent',
                    params={"parent": _subseries["uri"].rsplit("/")[-1], "position": 1},
                )
                return _file["uri"].rsplit("/")[-1]
            else:
                return _subseries["uri"].rsplit("/")[-1]
        elif "file" in kwargs.get("ancestors"):
            kwargs["level"] = "file"
            _file = _create_archival_object(
                asnake_client, test_name, resource_uri, **kwargs
            )
            return _file["uri"].rsplit("/")[-1]
        else:
            return None

    def _nest_archival_object(archival_object_uri, parent_id):
        asnake_client.post(
            f"{archival_object_uri}/parent", params={"parent": parent_id, "position": 1}
        )

    def _generate_audio_file(file_stem, **kwargs):
        audio = gTTS(text=file_stem.split("_")[-1])
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.mp3",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.mp3"
            )
        audio.save(output_file)
        print(f"🐞 GENERATED FILE {output_file}")

    def _generate_video_file(file_stem, **kwargs):
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.mp4",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.mp4"
            )
        timecode_text = r"text='timestamp:%{pts\:hms}': x=(w-text_w)/2: y=100: font=Mono: fontsize=64: fontcolor=Yellow: box=1: boxcolor=Black: boxborderw=10"
        testname_text = f'text={file_stem.split("_")[-1]}: x=(w-text_w)/2: y=(h-text_h)/2: fontsize=196: fontcolor=White: shadowx=3: shadowy=3'
        subprocess.run(
            [
                config("WORK_FFMPEG_CMD"),
                "-f",
                "lavfi",
                "-i",
                "smptebars=s=1280x720",
                "-vf",
                f"drawtext={timecode_text}, drawtext={testname_text}",
                "-c:v",
                "libx264",
                "-r",
                "30",
                "-frames:v",
                "90",
                "-y",
                output_file,
            ],
            check=True,
        )
        print(f"🐞 GENERATED FILE {output_file}")

    def _generate_image_file(file_stem, **kwargs):
        tmp_file, headers = urllib.request.urlretrieve(
            config("TEST_IMG_URL", default="https://picsum.photos/1600/1200")
        )
        img = Image.open(tmp_file)
        drw = ImageDraw.Draw(img)
        fnt = ImageFont.truetype(
            config(
                "TEST_IMG_FONT",
                default="/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            ),
            144,
        )
        drw.text(
            (800, 600),
            file_stem.split("_")[-1],
            fill="white",
            font=fnt,
            anchor="mm",
            stroke_width=10,
            stroke_fill="black",
        )
        if "directory" in kwargs:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"),
                kwargs["directory"],
                f"{file_stem}.tif",
            )
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
        else:
            output_file = os.path.join(
                config("INITIAL_ORIGINAL_FILES"), f"{file_stem}.tif"
            )
        img.save(output_file)
        img.close()
        print(f"🐞 GENERATED FILE {output_file}")

    def _create_digital_files(component_id, **kwargs):
        if "mixed" in component_id:
            # NOTE "mixed" will always generate one each of audio, video, image
            # regardless of archival_object_count value
            if component_id.endswith("1"):
                _generate_audio_file(component_id, **kwargs)
            elif component_id.endswith("2"):
                _generate_video_file(component_id, **kwargs)
            elif component_id.endswith("3"):
                _generate_image_file(component_id, **kwargs)
        elif "audio" in component_id:
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    _generate_audio_file(
                        "{}{}".format(
                            "".join(
                                random.choices(
                                    string.ascii_lowercase + string.digits, k=6
                                )
                            ),
                            str(i).zfill(2),
                        ),
                        directory=component_id,
                        **kwargs,
                    )
            else:
                _generate_audio_file(component_id, **kwargs)
        elif "video" in component_id:
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    _generate_video_file(
                        "{}{}".format(
                            "".join(
                                random.choices(
                                    string.ascii_lowercase + string.digits, k=6
                                )
                            ),
                            str(i).zfill(2),
                        ),
                        directory=component_id,
                        **kwargs,
                    )
            else:
                _generate_video_file(component_id, **kwargs)
        else:  # create images by default
            if kwargs.get("digital_file_count") > 1:
                for i in range(1, kwargs.get("digital_file_count") + 1):
                    _generate_image_file(
                        "{}{}".format(
                            "".join(
                                random.choices(
                                    string.ascii_lowercase + string.digits, k=6
                                )
                            ),
                            str(i).zfill(2),
                        ),
                        directory=component_id,
                        **kwargs,
                    )
            else:
                _generate_image_file(component_id, **kwargs)

    _content_attributes = []
    for j in range(1, kwargs.get("collection_count") + 1):
        _resource = _create_archivesspace_resource(asnake_client, test_name, j=j)
        parent_id = _create_ancestors(
            asnake_client,
            test_name,
            _resource["uri"],
            i=0,
            j=j,
            **kwargs,
        )
        for i in range(1, kwargs.get("archival_object_count") + 1):
            _archival_object = _create_archival_object(
                asnake_client,
                test_name,
                _resource["uri"],
                i=i,
                j=j,
                **kwargs,
            )
            if parent_id:
                _nest_archival_object(_archival_object["uri"], parent_id)
            _create_digital_files(_archival_object["component_id"], **kwargs)
            _content_attributes.append(
                {
                    "archival_object_uri": _archival_object["uri"],
                    "component_id": _archival_object["component_id"],
                    "resource_uri": _resource["uri"],
                    "id_0": _resource["id_0"],
                }
            )
    return _content_attributes


def update_content(page: Page, asnake_client, content_attributes, regenerate):
    if regenerate == "one":
        # VALIDATE EXISTING ALCHEMIST OBJECT
        alchemist_object_path = "/".join(
            [
                config("ALCHEMIST_URL_PREFIX"),
                content_attributes[0]["id_0"],
                content_attributes[0]["component_id"],
            ]
        )
        if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
            invalidate_cloudfront_path(
                path=f"/{alchemist_object_path}/*", caller_reference=str(time.time())
            )
        page.goto(
            "/".join([config("ALCHEMIST_BASE_URL").rstrip("/"), alchemist_object_path])
        )
        expect(page.locator("hgroup > h1")).not_to_contain_text("UPDATED")
        # SET VARIABLES
        archival_object_uri = content_attributes[0]["archival_object_uri"]
        # UPDATE ARCHIVAL OBJECT
        print(
            f'🐞 UPDATING ARCHIVAL OBJECT {content_attributes[0]["component_id"]}: {archival_object_uri}'
        )
        archival_object = asnake_client.get(archival_object_uri).json()
        archival_object["title"] = f"UPDATED {archival_object['title']}"
        # POST
        asnake_client.post(archival_object_uri, json=archival_object)
    elif regenerate == "collection":
        alchemist_collection_path = "/".join(
            [
                config("ALCHEMIST_URL_PREFIX"),
                content_attributes[0]["id_0"],
            ]
        )
        if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
            invalidate_cloudfront_path(
                path=f"/{alchemist_collection_path}/*",
                caller_reference=str(time.time()),
            )
        for content_attribute in content_attributes:
            # VALIDATE EXISTING ALCHEMIST OBJECTS
            page.goto(
                "/".join(
                    [
                        config("ALCHEMIST_BASE_URL").rstrip("/"),
                        alchemist_collection_path,
                        content_attributes[0]["component_id"],
                    ]
                )
            )
            expect(page.locator("hgroup > h1")).not_to_contain_text("UPDATED")
            if content_attribute["id_0"] == content_attributes[0]["id_0"]:
                # SET VARIABLES
                archival_object_uri = content_attribute["archival_object_uri"]
                # UPDATE ARCHIVAL OBJECT
                print(
                    f'🐞 UPDATING ARCHIVAL OBJECT {content_attribute["component_id"]}: {archival_object_uri}'
                )
                archival_object = asnake_client.get(archival_object_uri).json()
                archival_object["title"] = f"UPDATED {archival_object['title']}"
                # POST
                asnake_client.post(archival_object_uri, json=archival_object)
    elif regenerate == "all":
        if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
            invalidate_cloudfront_path(caller_reference=str(time.time()))
        for content_attribute in content_attributes:
            # VALIDATE EXISTING ALCHEMIST OBJECTS
            page.goto(
                "/".join(
                    [
                        config("ALCHEMIST_BASE_URL").rstrip("/"),
                        config("ALCHEMIST_URL_PREFIX"),
                        content_attribute["id_0"],
                        content_attribute["component_id"],
                    ]
                )
            )
            expect(page.locator("hgroup > h1")).not_to_contain_text("UPDATED")
            # SET VARIABLES
            archival_object_uri = content_attribute["archival_object_uri"]
            # UPDATE ARCHIVAL OBJECT
            print(
                f'🐞 UPDATING ARCHIVAL OBJECT {content_attribute["component_id"]}: {archival_object_uri}'
            )
            archival_object = asnake_client.get(archival_object_uri).json()
            archival_object["title"] = f"UPDATED {archival_object['title']}"
            # POST
            asnake_client.post(archival_object_uri, json=archival_object)


def delete_archivesspace_records(test_name, asnake_client, **kwargs):
    # DELETE ANY EXISTING TEST RECORDS IN ARCHIVESSPACE
    for j in range(1, kwargs.get("collection_count") + 1):
        print(f"🐞 DELETING ARCHIVESSPACE RECORDS FOR {test_name.split('_')[-1]}xx{j}")
        delete_archivesspace_test_records(
            asnake_client, f'{test_name.split("_")[-1]}xx{j}'
        )


def delete_filesystem_files():
    # DELETE ANY EXISTING TEST FILES IN FILESYSTEM
    for pathname in glob.glob(os.path.join(config("INITIAL_ORIGINAL_FILES"), "*")):
        os.system(f"/bin/rm -r {pathname}")
    for pathname in glob.glob(os.path.join(config("BATCH_SETS_DIRECTORY"), "*")):
        os.system(f"/bin/rm -r {pathname}")
    for pathname in glob.glob(
        os.path.join(
            config(
                "WORK_STILLAGE_FILES",
                default=os.path.join(
                    os.path.dirname(config("WORK_PRESERVATION_FILES")), ".STILLAGE"
                ),
            ),
            "*",
        )
    ):
        os.system(f"/bin/rm -r {pathname}")


def delete_content(test_name, asnake_client, s3_client, **kwargs):
    delete_archivesspace_records(test_name, asnake_client, **kwargs)
    delete_filesystem_files()
    # # DELETE INITIAL_ORIGINAL_FILES FROM FAILED TESTS
    # for pathname in glob.glob(os.path.join(config("INITIAL_ORIGINAL_FILES"), "*")):
    #     os.system(f"/bin/rm -r {pathname}")
    # # DELETE FILES IN BATCH_SETS_DIRECTORY FROM PASSED TESTS
    # for pathname in glob.glob(
    #     os.path.join(
    #         config("BATCH_SETS_DIRECTORY"), f'**/*{test_name.split("_")[-1]}*'
    #     ),
    #     recursive=True,
    # ):
    #     os.system(f"/bin/rm -r {pathname}")
    # if config("WORK_STILLAGE_FILES", default=""):
    #     os.system(f'/bin/rm -r {os.path.join(config("WORK_STILLAGE_FILES"), "*")}')
    # else:
    #     os.system(
    #         f'/bin/rm -r {os.path.join(os.path.dirname(config("WORK_PRESERVATION_FILES")), ".STILLAGE", "*")}'
    #     )
    # DELETE S3 OBJECTS
    if test_name.split("_")[1] == "alchemist":
        bucket = config("ALCHEMIST_BUCKET")
        prefix = f'{config("ALCHEMIST_URL_PREFIX")}/'
    elif test_name.split("_")[1] == "s3":
        bucket = config("PRESERVATION_BUCKET")
        prefix = f'{test_name.split("_")[-1]}'
    s3_response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
    if s3_response.get("Contents"):
        s3_keys = [{"Key": s3_object["Key"]} for s3_object in s3_response["Contents"]]
        print(f"🐞 DELETING S3 OBJECTS: {s3_keys}")
        s3_response = s3_client.delete_objects(
            Bucket=bucket, Delete={"Objects": s3_keys}
        )


@pytest.fixture
def run(page, asnake_client, s3_client, timestamp, request):
    # default parameters if not passed from calling test
    def _run(
        test_name,
        collection_count=1,
        archival_object_count=1,
        level="file",
        ancestors=["subseries", "series"],
        digital_file_count=1,
        expected_outcome="success",
        simulate_archivesspace_offline=False,
        regenerate=None,
        timeout=60000,
    ):
        print("...")
        # argument values are passed from the calling test
        unblock_archivesspace_ports()
        print("🐞 DELETING PREVIOUS TEST CONTENT")
        delete_content(
            test_name, asnake_client, s3_client, collection_count=collection_count
        )
        print("🐞 GENERATING TEST CONTENT")
        content_attributes = generate_content(
            test_name,
            asnake_client,
            collection_count=collection_count,
            archival_object_count=archival_object_count,
            level=level,
            ancestors=ancestors,
            digital_file_count=digital_file_count,
        )
        # RUN PROCESS
        if test_name.split("_")[1] == "alchemist":
            destinations = ["access"]
        elif test_name.split("_")[1] == "s3":
            destinations = ["cloud"]
        print("🐞 RUNNING DISTILLERY")
        run_distillery(page, destinations, outcome=expected_outcome, timeout=timeout)
        if regenerate:
            update_content(page, asnake_client, content_attributes, regenerate)
            print("🐞 RUNNING ALCHEMIST REGENERATE")
            run_alchemist_regenerate(
                page,
                content_attributes,
                regenerate=regenerate,
                simulate_archivesspace_offline=simulate_archivesspace_offline,
            )
        return content_attributes

    def _unblock_archivesspace_ports():
        unblock_archivesspace_ports()

    request.addfinalizer(_unblock_archivesspace_ports)

    return _run


def reset_files_and_records(test_name, asnake_client, s3_client, **kwargs):
    test_id = test_name.split("_")[-1]
    # TODO refactor file deletion
    for f in glob.glob(os.path.join(config("INITIAL_ORIGINAL_FILES"), "*")):
        shutil.move(f, config("TEST_FILES", default="tests/files"))
    if "mixed" in test_name:
        # TODO 🐞 debug hang in magick stream
        generate_image_file(test_name, directory=test_name)
        generate_video_file(test_name, directory=test_name)
        generate_audio_file(test_name, directory=test_name)
        # TODO refactor file deletion
        os.system(
            f'/bin/rm -r {os.path.join(config("TEST_FILES", default="tests/files"), f"item-{test_name}")}'
        )
        if config("WORK_STILLAGE_FILES", default=""):
            os.system(f'/bin/rm -r {os.path.join(config("WORK_STILLAGE_FILES"), "*")}')
        else:
            os.system(
                f'/bin/rm -r {os.path.join(os.path.dirname(config("WORK_PRESERVATION_FILES")), ".STILLAGE", "*")}'
            )
    elif "image" in test_name:
        if "file_count" in kwargs and kwargs["file_count"] > 1:
            for i in range(kwargs["file_count"]):
                generate_image_file(
                    "{}{}".format(
                        "".join(
                            random.choices(string.ascii_lowercase + string.digits, k=6)
                        ),
                        str(i).zfill(2),
                    ),
                    directory=test_name,
                )
        else:
            generate_image_file(test_name)
        # TODO refactor file deletion
        os.system(
            f'/bin/rm -r {os.path.join(config("TEST_FILES", default="tests/files"), f"*{test_id}*")}'
        )
    elif "video" in test_name:
        generate_video_file(test_name)
        # TODO refactor file deletion
        os.system(
            f'/bin/rm -r {os.path.join(config("TEST_FILES", default="tests/files"), f"item-{test_name}.mp4")}'
        )
    elif "audio" in test_name:
        generate_audio_file(test_name)
        # TODO refactor file deletion
        os.system(
            f'/bin/rm -r {os.path.join(config("TEST_FILES", default="tests/files"), f"item-{test_name}.mp3")}'
        )
    else:
        move_test_files_to_initial_original_files_directory(test_name)
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # DELETE S3 OBJECTS
    if test_name.split("_")[1] == "s3":
        delete_s3_preservation_objects(s3_client, test_id)


def generate_records(test_name, asnake_client, **kwargs):
    """TODO expand as needed for custom kwargs"""
    test_id = test_name.split("_")[-1]
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    # NOTE kwargs["item"] should be None or dict with customized values
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client,
        test_name,
        test_id,
        resource_create_response.json()["uri"],
        customizations=kwargs.get("item"),
    )
    print(f"🐞 item_create_response:{test_id}", item_create_response.json())
    # NOTE kwargs["series"] should be None or dict with customized values
    if kwargs.get("series") is None:
        return {
            "resource_create_response": resource_create_response,
            "item_create_response": item_create_response,
        }
    else:
        # CREATE ARCHIVAL OBJECT SERIES RECORD
        series_create_response = create_archivesspace_test_archival_object_series(
            asnake_client,
            test_id,
            resource_create_response.json()["uri"],
            customizations=kwargs["series"],
        )
        print(f"🐞 series_create_response:{test_id}", series_create_response.json())
        # NOTE kwargs["subseries"] should be None or dict with customized values
        if kwargs.get("subseries") is None:
            # set item as a child of series
            item_parent_position_post_response = asnake_client.post(
                f'{item_create_response.json()["uri"]}/parent',
                params={
                    "parent": subseries_create_response.json()["id"],
                    "position": 1,
                },
            )
            print(
                "🐞 item_parent_position_post_response",
                item_parent_position_post_response.json(),
            )
            return {
                "resource_create_response": resource_create_response,
                "series_create_response": series_create_response,
                "item_create_response": item_create_response,
            }
        else:
            # CREATE ARCHIVAL OBJECT SUBSERIES RECORD
            subseries_create_response = (
                create_archivesspace_test_archival_object_subseries(
                    asnake_client,
                    test_id,
                    resource_create_response.json()["uri"],
                    customizations=kwargs["subseries"],
                )
            )
            print(
                f"🐞 subseries_create_response:{test_id}",
                subseries_create_response.json(),
            )
            # set subseries as a child of series
            subseries_parent_position_post_response = asnake_client.post(
                f'{subseries_create_response.json()["uri"]}/parent',
                params={"parent": series_create_response.json()["id"], "position": 1},
            )
            print(
                "🐞 subseries_parent_position_post_response",
                subseries_parent_position_post_response.json(),
            )
            # set item as a child of subseries
            item_parent_position_post_response = asnake_client.post(
                f'{item_create_response.json()["uri"]}/parent',
                params={
                    "parent": subseries_create_response.json()["id"],
                    "position": 1,
                },
            )
            print(
                "🐞 item_parent_position_post_response",
                item_parent_position_post_response.json(),
            )
            return {
                "resource_create_response": resource_create_response,
                "series_create_response": series_create_response,
                "subseries_create_response": subseries_create_response,
                "item_create_response": item_create_response,
            }


@pytest.fixture
def setup_test(page, asnake_client, s3_client, timestamp):
    def _setup_test(
        test_name,
        series=None,
        subseries=None,
        item=None,
        file_count=1,
        outcome="success",
    ):
        print("🐞 _setup_test...")
        print(f"🐞 series: {series}")
        print(f"🐞 subseries: {subseries}")
        print(f"🐞 file_count: {file_count}")
        print(f"🐞 outcome: {outcome}")
        reset_files_and_records(
            test_name, asnake_client, s3_client, file_count=file_count
        )
        generated_records_responses = generate_records(
            test_name, asnake_client, series=series, subseries=subseries, item=item
        )
        # RUN PROCESS
        if test_name.split("_")[1] == "alchemist":
            destinations = ["access"]
        elif test_name.split("_")[1] == "s3":
            destinations = ["cloud"]
        run_distillery(page, destinations, outcome=outcome)
        alchemist_item_uri = format_alchemist_item_uri(
            test_name, test_name.split("_")[-1]
        )
        # INVALIDATE CLOUDFRONT ITEMS
        if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
            invalidate_cloudfront_path(caller_reference=timestamp)
        return {
            **generated_records_responses,
            "alchemist_item_uri": alchemist_item_uri,
        }

    return _setup_test


@pytest.mark.skipif(
    not os.getenv("DELETE_ARCHIVESSPACE_TEST_RECORDS"),
    reason="environment variable DELETE_ARCHIVESSPACE_TEST_RECORDS is not set",
)
def test_delete_archivesspace_test_records(asnake_client):
    test_identifiers = [
        name.rsplit("_", maxsplit=1)[-1]
        for name, obj in inspect.getmembers(sys.modules[__name__])
        if (
            inspect.isfunction(obj)
            and name.startswith("test_")
            and name != "test_delete_archivesspace_test_records"
        )
    ]
    print("🐞 DELETING TEST RECORDS")
    for test_id in test_identifiers:
        print(f"🐞 {test_id}")
        delete_archivesspace_test_records(asnake_client, test_id)


def test_distillery_landing(page: Page):
    page.goto(config("DISTILLERY_BASE_URL"))
    expect(page).to_have_title("Distillery")


def test_alchemist_fail_unpublished_archival_object_sjex6(page: Page, asnake_client):
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(test_name)
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client, test_name, test_id, resource_create_response.json()["uri"]
    )
    print(f"🐞 archival_object_create_response:{test_id}", item_create_response.json())
    # CUSTOMIZE ARCHIVAL OBJECT ITEM RECORD
    # set publish to False
    item = asnake_client.get(item_create_response.json()["uri"]).json()
    item["publish"] = False
    item_update_response = asnake_client.post(item["uri"], json=item)
    print(f"🐞 item_update_response:{test_id}", item_update_response.json())
    # RUN DISTILLERY ACCESS WORKFLOW
    page = run_distillery(page, ["access"], outcome="failure")
    # TODO check contents of iframe


def test_alchemist_fail_unpublished_ancestor_jvycv(page: Page, asnake_client):
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(test_name)
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CUSTOMIZE RESOURCE RECORD
    # set publish to False
    resource = asnake_client.get(resource_create_response.json()["uri"]).json()
    resource["publish"] = False
    resource_update_response = asnake_client.post(resource["uri"], json=resource)
    print(f"🐞 resource_update_response:{test_id}", resource_update_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client, test_name, test_id, resource_create_response.json()["uri"]
    )
    print(f"🐞 item_create_response:{test_id}", item_create_response.json())
    # RUN DISTILLERY ACCESS WORKFLOW
    page = run_distillery(page, ["access"], outcome="failure")
    # TODO check contents of iframe


def test_alchemist_archivesspace_fail_file_version_exists_70811b(page: Page):
    """Fail validation when file_version exists."""
    test_name = inspect.currentframe().f_code.co_name
    # BEGIN BASIC ALCHEMIST PROCESS
    kwargs = {
        "collection_count": 1,
        "archival_object_count": 1,
        "level": "file",
        "ancestors": [],
        "digital_file_count": 1,
    }
    # kwargs: collection_count
    delete_archivesspace_records(test_name, asnake_client, **kwargs)
    delete_filesystem_files()
    # kwargs: collection_count, archival_object_count, level, ancestors
    generate_archivesspace_records(test_name, asnake_client, **kwargs)
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    generate_filesystem_files(test_name, **kwargs)
    # kwargs: file_versions_op, thumbnail_label, outcome, timeout
    publish_to_web(page, **kwargs)
    # END BASIC ALCHEMIST PROCESS
    # BEGIN ATTEMPT TO PUBLISH AN EXISTING OBJECT
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    generate_filesystem_files(test_name, **kwargs)
    # kwargs: file_versions_op, thumbnail_label, outcome, timeout
    publish_to_web(page, outcome="failure", **kwargs)
    # END ATTEMPT TO PUBLISH AN EXISTING OBJECT


def test_alchemist_archivesspace_replace_file_version_exists_328abf(
    page: Page, asnake_client
):
    """Replace records and files when file_version exists."""
    test_name = inspect.currentframe().f_code.co_name
    # BEGIN BASIC ALCHEMIST PROCESS
    kwargs = {
        "collection_count": 1,
        "archival_object_count": 1,
        "level": "file",
        "ancestors": [],
        "digital_file_count": 1,
    }
    # kwargs: collection_count
    delete_archivesspace_records(test_name, asnake_client, **kwargs)
    delete_filesystem_files()
    # kwargs: collection_count, archival_object_count, level, ancestors
    record_data = generate_archivesspace_records(test_name, asnake_client, **kwargs)
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    generate_filesystem_files(test_name, **kwargs)
    # kwargs: file_versions_op, thumbnail_label, outcome, timeout
    publish_to_web(page, **kwargs)
    # END BASIC ALCHEMIST PROCESS
    # BEGIN ATTEMPT TO PUBLISH AN EXISTING OBJECT
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    generate_filesystem_files(test_name, **kwargs)
    # kwargs: file_versions_op, thumbnail_label, outcome, timeout
    publish_to_web(page, file_versions_op="replace", **kwargs)
    # END ATTEMPT TO PUBLISH AN EXISTING OBJECT
    for _ in record_data:
        archival_object = asnake_client.get(
            f'{_["archival_object_uri"]}?resolve[]=digital_object'
        ).json()
        # create_time of digital_object should be different from file_version
        # print(f'🐞 {archival_object["instances"][0]["digital_object"]["_resolved"]["create_time"]}')
        # print(f'🐞 {archival_object["instances"][0]["digital_object"]["_resolved"]["system_mtime"]}')
        # print(f'🐞 {archival_object["instances"][0]["digital_object"]["_resolved"]["user_mtime"]}')
        # print(f'🐞 {archival_object["instances"][0]["digital_object"]["_resolved"]["file_versions"][0]["create_time"]}')
        # print(f'🐞 {archival_object["instances"][0]["digital_object"]["_resolved"]["file_versions"][0]["system_mtime"]}')
        # print(f'🐞 {archival_object["instances"][0]["digital_object"]["_resolved"]["file_versions"][0]["user_mtime"]}')
        assert (
            len(
                {
                    archival_object["instances"][0]["digital_object"]["_resolved"][
                        "create_time"
                    ],
                    archival_object["instances"][0]["digital_object"]["_resolved"][
                        "file_versions"
                    ][0]["create_time"],
                }
            )
            > 1
        )


def test_alchemist_archivesspace_file_uri_39d132(run, page: Page, asnake_client):
    """Confirm ArchivesSpace file_uri matches Alchemist item URL."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        level="item",
        ancestors=[],
    )
    # VALIDATE URI/URL MATCH
    for _ in run_output:
        alchemist_item_url = "/".join(
            [
                config("ALCHEMIST_BASE_URL").rstrip("/"),
                config("ALCHEMIST_URL_PREFIX"),
                _["id_0"],
                _["component_id"],
            ]
        )
        print(f"🐞 {alchemist_item_url}")
        page.goto(alchemist_item_url)
        archival_object = asnake_client.get(
            f'{_["archival_object_uri"]}?resolve[]=digital_object'
        ).json()
        assert len(archival_object["instances"]) == 1
        assert (
            archival_object["instances"][0]["digital_object"]["_resolved"][
                "file_versions"
            ][1]["file_uri"]
            == alchemist_item_url
        )
        assert (
            archival_object["instances"][0]["digital_object"]["_resolved"]["publish"]
            is True
        )


def test_alchemist_thumbnaillabel_sequence_yw3ff(page: Page, asnake_client, timestamp):
    """Use sequence indicator as label in Universal Viewer."""
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(test_name)
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client, test_name, test_id, resource_create_response.json()["uri"]
    )
    print(f"🐞 item_create_response:{test_id}", item_create_response.json())
    # RUN ALCHEMIST PROCESS
    run_distillery(page, ["access"])
    alchemist_item_uri = format_alchemist_item_uri(test_name, test_id)
    # INVALIDATE CLOUDFRONT ITEMS
    if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
        invalidate_cloudfront_path(caller_reference=timestamp)
    # VALIDATE ALCHEMIST ITEM
    page.goto(alchemist_item_uri)
    expect(page.locator("#thumb-0")).to_have_text("1")
    expect(page.locator("#thumb-1")).to_have_text("2")
    expect(page.locator("#thumb-2")).to_have_text("last")


def test_alchemist_thumbnaillabel_filename_wef99(page: Page, asnake_client, timestamp):
    """Use filename as label in Universal Viewer."""
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(test_name)
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client, test_name, test_id, resource_create_response.json()["uri"]
    )
    print(f"🐞 item_create_response:{test_id}", item_create_response.json())
    # RUN ALCHEMIST PROCESS
    run_distillery(page, ["access"], thumbnail_label="filename")
    alchemist_item_uri = format_alchemist_item_uri(test_name, test_id)
    # INVALIDATE CLOUDFRONT ITEMS
    if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
        invalidate_cloudfront_path(caller_reference=timestamp)
    # VALIDATE ALCHEMIST ITEM
    page.goto(alchemist_item_uri)
    expect(page.locator("#thumb-0")).to_have_text("lQGJCMY5qcM-unsplash_001")
    expect(page.locator("#thumb-1")).to_have_text("lQGJCMY5qcM-unsplash_002")
    expect(page.locator("#thumb-2")).to_have_text("lQGJCMY5qcM-unsplash_last")


def test_alchemist_regenerate_one_986204(run, page: Page, asnake_client, timestamp):
    """Regenerate all Alchemist assets for one object."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        archival_object_count=2,
        level="file",
        ancestors=["subseries", "series"],
        simulate_archivesspace_offline=True,
        regenerate="one",
    )
    # VALIDATE REGENERATED ALCHEMIST PAGES
    for _ in run_output:
        page.goto(
            "/".join(
                [
                    config("ALCHEMIST_BASE_URL").rstrip("/"),
                    config("ALCHEMIST_URL_PREFIX"),
                    _["id_0"],
                    _["component_id"],
                ]
            )
        )
        if _["component_id"] == run_output[0]["component_id"]:
            expect(page.locator("hgroup > h1")).to_contain_text("UPDATED")
        else:
            expect(page.locator("hgroup > h1")).not_to_contain_text("UPDATED")


def test_alchemist_regenerate_collection_84bc89(
    run, page: Page, asnake_client, timestamp
):
    """Regenerate all Alchemist assets for a collection."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        collection_count=2,
        level="file",
        ancestors=["subseries", "series"],
        simulate_archivesspace_offline=True,
        regenerate="collection",
    )
    # VALIDATE REGENERATED ALCHEMIST PAGES
    for _ in run_output:
        page.goto(
            "/".join(
                [
                    config("ALCHEMIST_BASE_URL").rstrip("/"),
                    config("ALCHEMIST_URL_PREFIX"),
                    _["id_0"],
                    _["component_id"],
                ]
            )
        )
        if _["id_0"] == run_output[0]["id_0"]:
            expect(page.locator("hgroup > h1")).to_contain_text("UPDATED")
        else:
            expect(page.locator("hgroup > h1")).not_to_contain_text("UPDATED")


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_alchemist_regenerate_collection_84bc89xx1():
    return


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_alchemist_regenerate_collection_84bc89xx2():
    return


def test_alchemist_regenerate_all_546ebc(run, page: Page, asnake_client, timestamp):
    """Regenerate all Alchemist assets."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        collection_count=2,
        level="file",
        ancestors=["subseries", "series"],
        simulate_archivesspace_offline=True,
        regenerate="all",
    )
    # VALIDATE REGENERATED ALCHEMIST PAGES
    for _ in run_output:
        page.goto(
            "/".join(
                [
                    config("ALCHEMIST_BASE_URL").rstrip("/"),
                    config("ALCHEMIST_URL_PREFIX"),
                    _["id_0"],
                    _["component_id"],
                ]
            )
        )
        expect(page.locator("hgroup > h1")).to_contain_text("UPDATED")


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_alchemist_regenerate_all_546ebcxx1():
    return


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_alchemist_regenerate_all_546ebcxx2():
    return


def test_alchemist_fileversions_fail_2tgwm(page: Page, asnake_client):
    """Fail validation when digital_object file_versions exist."""
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(test_name)
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client, test_name, test_id, resource_create_response.json()["uri"]
    )
    print(f"🐞 item_create_response:{test_id}", item_create_response.json())
    # CREATE DIGITAL OBJECT RECORD
    (
        digital_object_create_response,
        digital_object_id,
    ) = create_archivesspace_test_digital_object(asnake_client, test_name, test_id)
    print(
        f"🐞 digital_object_create_response:{test_id}",
        digital_object_create_response.json(),
    )
    # UPDATE ARCHIVAL OBJECT ITEM RECORD
    item = asnake_client.get(item_create_response.json()["uri"]).json()
    # update title
    item["instances"] = [
        {
            "instance_type": "digital_object",
            "digital_object": {"ref": digital_object_create_response.json()["uri"]},
        }
    ]
    item_update_response = asnake_client.post(item["uri"], json=item)
    print(f"🐞 item_update_response:{test_id}", item_update_response.json())
    # RUN ALCHEMIST PROCESS
    run_distillery(page, ["access"], outcome="failure")
    # TODO check contents of iframe


def test_alchemist_item_breadcrumbs_multiple_single_image_objects_edcb48(
    run, page: Page, asnake_client
):
    """Publish multiple single-image archival objects and check breadcrumbs."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        collection_count=2,
        archival_object_count=2,
        level="item",
        ancestors=["file", "subseries", "series"],
        timeout=120000,
    )
    # VALIDATE ALCHEMIST DISPLAY
    for i in run_output:
        alchemist_item_url = "/".join(
            [
                config("ALCHEMIST_BASE_URL").rstrip("/"),
                config("ALCHEMIST_URL_PREFIX"),
                i["id_0"],
                i["component_id"],
            ]
        )
        print(f"🐞 {alchemist_item_url}")
        page.goto(alchemist_item_url)
        # validate breadcrumbs
        expect(page.locator("hgroup nav li:nth-child(1)")).to_have_text(
            f'{test_name.capitalize().rsplit("_", maxsplit=1)[0].replace("_", " ")} {i["id_0"].replace("xx", " ")}'
        )
        expect(page.locator("hgroup nav li:nth-child(2)")).to_contain_text(
            f'series {i["id_0"].split("x")[0]}'
        )
        expect(page.locator("hgroup nav li:nth-child(3)")).to_contain_text(
            f'subseries {i["id_0"].split("x")[0]}'
        )
        expect(page.locator("hgroup nav li:nth-child(4)")).to_contain_text(
            f'file {i["id_0"].split("x")[0]}'
        )
        expect(page.locator("hgroup nav li:nth-child(5)")).to_have_text(
            f'open the {i["component_id"].split("_")[0]} {i["component_id"].split("_")[-1].replace("xx", " ")} collection guide metadata record'
        )


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_alchemist_item_breadcrumbs_multiple_single_image_objects_edcb48xx1():
    return


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_alchemist_item_breadcrumbs_multiple_single_image_objects_edcb48xx2():
    return


def test_alchemist_file_breadcrumbs_multi_image_object_67707b(
    run, page: Page, asnake_client
):
    """Publish a multi-image archival object and check breadcrumbs."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        digital_file_count=2,
    )
    # VALIDATE ALCHEMIST DISPLAY
    for i in run_output:
        alchemist_item_url = "/".join(
            [
                config("ALCHEMIST_BASE_URL").rstrip("/"),
                config("ALCHEMIST_URL_PREFIX"),
                i["id_0"],
                i["component_id"],
            ]
        )
        print(f"🐞 {alchemist_item_url}")
        page.goto(alchemist_item_url)
        # validate breadcrumbs
        expect(page.locator("hgroup nav li:nth-child(1)")).to_have_text(
            f'{test_name.capitalize().rsplit("_", maxsplit=1)[0].replace("_", " ")} {i["id_0"].replace("xx", " ")}'
        )
        expect(page.locator("hgroup nav li:nth-child(2)")).to_contain_text(
            f'series {i["id_0"].split("x")[0]}'
        )
        expect(page.locator("hgroup nav li:nth-child(3)")).to_contain_text(
            f'subseries {i["id_0"].split("x")[0]}'
        )
        expect(page.locator("hgroup nav li:nth-child(4)")).to_have_text(
            f'open the {i["component_id"].split("_")[0]} {i["component_id"].split("_")[-1].replace("xx", " ")} collection guide metadata record'
        )


def test_alchemist_singleitem_video_frq8s(page: Page, asnake_client, timestamp):
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(
        "test_alchemist_singleitem_video_frq8s.mp4"
    )
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client, test_name, test_id, resource_create_response.json()["uri"]
    )
    print(f"🐞 item_create_response:{test_id}", item_create_response.json())
    # RUN ALCHEMIST PROCESS
    run_distillery(page, ["access"])
    alchemist_item_uri = format_alchemist_item_uri(test_name, test_id)
    # INVALIDATE CLOUDFRONT ITEMS
    if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
        invalidate_cloudfront_path(caller_reference=timestamp)
    # VALIDATE ALCHEMIST HTML
    page.goto(alchemist_item_uri)
    expect(page.locator("video")).to_be_visible()


def test_alchemist_kitchen_sink_multi_item_image_k76zs(
    setup_test, page: Page, asnake_client
):
    """Display every metadata field with a multi-item image object."""
    test_name = inspect.currentframe().f_code.co_name
    # CUSTOMIZE ARCHIVAL OBJECT SERIES RECORD
    series = {}
    series["title"] = "Series Title Longer Than 50 Characters"
    # CUSTOMIZE ARCHIVAL OBJECT SUBSERIES RECORD
    subseries = {}
    subseries["title"] = "Subseries Title Longer Than 50 Characters"
    # CUSTOMIZE AGENT PERSON RECORDS
    linked_agents = [
        {
            "role": "creator",
        },
        {
            "role": "subject",
        },
        {
            "relator": "ard",
            "role": "creator",
            "publish": True,
        },
        {
            "relator": "act",
            "role": "subject",
            "publish": True,
        },
    ]
    # CUSTOMIZE ARCHIVAL OBJECT ITEM RECORD
    item = {}
    # add dates
    item["dates"] = [
        {"label": "digitized", "date_type": "single", "begin": "2022-02-22"},
        {"label": "creation", "date_type": "single", "begin": "1584-02-29"},
        {"label": "creation", "date_type": "inclusive", "begin": "1900", "end": "1901"},
        {
            "label": "creation",
            "date_type": "inclusive",
            "begin": "1911-01",
            "end": "1911-12",
        },
        {
            "label": "creation",
            "date_type": "inclusive",
            "begin": "1969-12-31",
            "end": "1970-01-01",
        },
        {
            "label": "creation",
            "date_type": "bulk",
            "begin": "1999-12-31",
            "end": "2000-01-01",
        },
        {
            "label": "creation",
            "date_type": "single",
            "expression": "ongoing into the future",
        },
    ]
    # add linked_agents
    item["linked_agents"] = linked_agents
    # add extents
    item["extents"] = [
        {"portion": "whole", "number": "1", "extent_type": "books"},
        {"portion": "part", "number": "2", "extent_type": "photographs"},
    ]
    # add subjects
    item["subjects"] = [{"ref": "/subjects/1"}, {"ref": "/subjects/2"}]
    # add notes
    item["notes"] = [
        {
            "jsonmodel_type": "note_singlepart",
            "type": "abstract",
            "content": [
                "Published note. One content item. Sint nulla ea nostrud est tempor non exercitation tempor ad consectetur nisi voluptate consequat."
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_singlepart",
            "type": "materialspec",
            "content": [
                "Published note. Multiple content items: One. Veniam enim ullamco non commodo enim ad incididunt quis.",
                "Published note. Multiple content items: Two. Enim tempor ea nulla voluptate incididunt voluptate.",
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_singlepart",
            "type": "abstract",
            "label": "Foo Note",
            "content": [
                "Published note. One content item. Laborum labore irure consequat dolore aute minim deserunt nostrud amet."
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_singlepart",
            "type": "abstract",
            "label": "Foo Note",
            "content": [
                "Published note. Multiple content items: One. Consequat cupidatat enim duis Lorem ipsum.",
                "Published note. Multiple content items: Two. Lorem ipsum velit cillum ex do officia pariatur pariatur duis est dolor.",
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_singlepart",
            "type": "abstract",
            "content": [
                "Unpublished note. One content item. Enim aute Lorem tempor exercitation enim adipisicing occaecat veniam ad duis excepteur culpa ut consectetur."
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_singlepart",
            "type": "abstract",
            "content": [
                "Unpublished note. Multiple content items: One. Consectetur cupidatat ea sunt sit enim minim officia ea ut tempor aute.",
                "Unpublished note. Multiple content items: Two. Cillum dolore amet dolor labore do deserunt adipisicing dolore in aliquip nulla.",
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. One published text subnote. Consequat nostrud ipsum irure reprehenderit qui veniam pariatur.",
                    "publish": True,
                }
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "userestrict",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. Multiple published text subnotes: One. Officia nisi nisi incididunt excepteur nisi.",
                    "publish": True,
                },
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. Multiple published text subnotes: Two. Laborum commodo exercitation deserunt velit.",
                    "publish": True,
                },
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. One published text subnote. Laboris ipsum cupidatat consequat velit.",
                    "publish": True,
                }
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. Multiple published text subnotes: One. Laborum anim laborum est laborum.",
                    "publish": True,
                },
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. Multiple published text subnotes: Two. Consectetur laborum laborum quis.",
                    "publish": True,
                },
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. One unpublished text subnote. Laborum cupidatat adipisicing cillum deserunt.",
                    "publish": False,
                }
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "label": "Baz Note",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. One unpublished text subnote. Aliqua consequat mollit reprehenderit pariatur exercitation nisi culpa incididunt.",
                    "publish": False,
                }
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. Multiple unpublished text subnotes: One. Aliquip culpa pariatur consequat.",
                    "publish": False,
                },
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. Multiple unpublished text subnotes: Two. Tempor exercitation sunt.",
                    "publish": False,
                },
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. One unpublished text subnote. Esse in proident.",
                    "publish": False,
                }
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. Multiple unpublished text subnotes: One. Laborum laborum sunt.",
                    "publish": False,
                },
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. Multiple unpublished text subnotes: Two. Sint adipisicing.",
                    "publish": False,
                },
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "label": "Bar Note",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. Mixed publication status text subnotes: Published. Laboris id cupidatat.",
                    "publish": True,
                },
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. Mixed publication status text subnotes: Unpublished. Nostrud anim dolore anim consequat quis sit laborum non.",
                    "publish": False,
                },
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. Mixed publication status text subnotes: Published. Est nostrud laboris id sint amet proident officia commodo ut sint amet sint dolore sunt.",
                    "publish": True,
                },
                {
                    "jsonmodel_type": "note_text",
                    "content": "Unpublished note. Mixed publication status text subnotes: Unpublished. Fugiat irure sunt magna nulla minim commodo dolor ea dolor aliquip enim magna fugiat.",
                    "publish": False,
                },
            ],
            "publish": False,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "label": "Bar Note",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. One published text subnote. Minim aute nulla laborum ullamco do incididunt nostrud irure eiusmod laborum elit deserunt.",
                    "publish": True,
                }
            ],
            "publish": True,
        },
        {
            "jsonmodel_type": "note_multipart",
            "type": "scopecontent",
            "label": "Foo Note",
            "subnotes": [
                {
                    "jsonmodel_type": "note_text",
                    "content": "Published note. One published text subnote. Eiusmod irure laboris eu reprehenderit proident exercitation qui nulla irure amet.",
                    "publish": True,
                }
            ],
            "publish": True,
        },
    ]
    setup_output = setup_test(
        test_name, series=series, subseries=subseries, item=item, file_count=2
    )
    # VALIDATE ALCHEMIST DISPLAY
    page.goto(setup_output["alchemist_item_uri"])
    expect(page.locator("hgroup p:first-of-type")).to_have_text(
        "1584 February 29; 1900 to 1901; 1911 January to December; 1969 December 31 to 1970 January 1; 1999 December 31 to 2000 January 1; ongoing into the future"
    )
    expect(page.locator("hgroup nav li:nth-child(1)")).to_have_text(
        f'{test_name.capitalize().replace("_", " ")}'
    )
    expect(page.locator("hgroup nav li:nth-child(2)")).to_have_text(
        f"Series Title Longer Than 50 Characters"
    )
    expect(page.locator("hgroup nav li:nth-child(3)")).to_have_text(
        f"Subseries Title Longer Than 50 Characters"
    )
    expect(page.locator("hgroup nav li:nth-child(4)")).to_have_text(
        f'open the Item {test_name.split("_")[-1]} collection guide metadata record'
    )
    expect(page.locator("#metadata")).to_contain_text("Collection")
    expect(page.locator("#metadata")).to_contain_text("Series")
    expect(page.locator("#metadata")).to_contain_text("Sub-Series")
    expect(page.locator("#metadata dt").filter(has_text="Identifier")).to_be_visible()
    expect(
        page.get_by_role(
            "link",
            name=f'{test_name.split("_")[-1].capitalize()}, Published Creator [Artistic director]',
        )
    ).to_be_visible()
    expect(page.get_by_role("link", name="Commencement")).to_be_visible()
    expect(page.get_by_role("link", name="Conferences")).to_be_visible()
    expect(
        page.get_by_role(
            "link",
            name=f'{test_name.split("_")[-1].capitalize()}, Published Subject [Actor]',
        )
    ).to_be_visible()
    expect(page.locator("#metadata")).to_contain_text("1 books", ignore_case=True)
    expect(page.locator("#metadata")).to_contain_text("2 photographs", ignore_case=True)
    expect(page.locator("#metadata")).to_contain_text("Abstract")
    expect(page.locator("#metadata")).to_contain_text("Materials Specific Details")
    expect(page.locator("#metadata")).to_contain_text("Foo Note")
    expect(page.locator("#metadata")).to_contain_text("Scope and Contents")
    expect(page.locator("#metadata")).to_contain_text("Conditions Governing Use")
    expect(page.locator("#metadata")).to_contain_text("Bar Note")
    expect(page.locator("#metadata")).not_to_contain_text(
        "unpublished", ignore_case=True
    )


def test_alchemist_archivesspace_digital_object_type_image_pw83h(
    setup_test, asnake_client
):
    """Confirm digital_object_type is set to still_image."""
    test_name = inspect.currentframe().f_code.co_name
    setup_output = setup_test(test_name)
    # VALIDATE DIGITAL OBJECT RECORD
    archival_object = asnake_client.get(
        f'{setup_output["item_create_response"].json()["uri"]}?resolve[]=digital_object'
    ).json()
    assert (
        archival_object["instances"][0]["digital_object"]["_resolved"][
            "digital_object_type"
        ]
        == "still_image"
    )


def test_alchemist_archivesspace_digital_object_type_video_zewtg(
    setup_test, asnake_client
):
    """Confirm digital_object_type is set to moving_image."""
    test_name = inspect.currentframe().f_code.co_name
    setup_output = setup_test(test_name)
    # VALIDATE DIGITAL OBJECT RECORD
    archival_object = asnake_client.get(
        f'{setup_output["item_create_response"].json()["uri"]}?resolve[]=digital_object'
    ).json()
    assert (
        archival_object["instances"][0]["digital_object"]["_resolved"][
            "digital_object_type"
        ]
        == "moving_image"
    )


def test_alchemist_archivesspace_digital_object_type_audio_jwkr2(
    setup_test, asnake_client
):
    """Confirm digital_object_type is set to sound_recording."""
    test_name = inspect.currentframe().f_code.co_name
    # TODO update outcome when alchemist handles audio
    setup_output = setup_test(test_name, outcome="failure")
    # # VALIDATE DIGITAL OBJECT RECORD
    # archival_object = asnake_client.get(
    #     f'{setup_output["item_create_response"].json()["uri"]}?resolve[]=digital_object'
    # ).json()
    # assert (
    #     archival_object["instances"][0]["digital_object"]["_resolved"][
    #         "digital_object_type"
    #     ]
    #     == "sound_recording"
    # )


def test_s3_archivesspace_entries_multiple_single_image_objects_7d6f50(
    run, asnake_client, s3_client
):
    """Ensure single-image archival objects have entries in S3 and ArchivesSpace."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        collection_count=2,
        archival_object_count=2,
        level="item",
        ancestors=[],
        timeout=120000,
    )
    # VALIDATE S3 UPLOADS AND ARCHIVESSPACE RECORDS
    for i in run_output:
        # get a list of s3 objects under this collection prefix
        s3_response = s3_client.list_objects_v2(
            Bucket=config("PRESERVATION_BUCKET"), Prefix=i["id_0"]
        )
        # get the digital_object tree
        archival_object = asnake_client.get(
            f'{i["archival_object_uri"]}?resolve[]=digital_object'
        ).json()
        assert len(archival_object["instances"]) == 1
        tree = asnake_client.get(
            f'{archival_object["instances"][0]["digital_object"]["ref"]}/tree/root'
        ).json()
        for waypoint in tree["precomputed_waypoints"][""]["0"]:
            # ensure the s3 key matches the value in the file_uri_summary
            assert waypoint["file_uri_summary"].split(
                f's3://{config("PRESERVATION_BUCKET")}/'
            )[-1] in [s3_object["Key"] for s3_object in s3_response["Contents"]]
            # ensure the original filename from the s3 key matches the digital_object_component label
            assert waypoint["label"] in [
                s3_object["Key"].split("/")[-2] for s3_object in s3_response["Contents"]
            ]


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_s3_archivesspace_entries_multiple_single_image_objects_7d6f50xx1():
    return


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_s3_archivesspace_entries_multiple_single_image_objects_7d6f50xx2():
    return


def test_s3_archivesspace_entries_multiple_multi_image_objects_3355ff(
    run, asnake_client, s3_client
):
    """Ensure multi-image archival objects have entries in S3 and ArchivesSpace."""
    test_name = inspect.currentframe().f_code.co_name
    run_output = run(
        test_name,
        collection_count=2,
        archival_object_count=2,
        level="file",
        ancestors=["subseries", "series"],
        digital_file_count=2,
        timeout=120000,
    )
    # VALIDATE S3 UPLOADS AND ARCHIVESSPACE RECORDS
    for i in run_output:
        # get a list of s3 objects under this collection prefix
        s3_response = s3_client.list_objects_v2(
            Bucket=config("PRESERVATION_BUCKET"), Prefix=i["id_0"]
        )
        # get the digital_object tree
        archival_object = asnake_client.get(
            f'{i["archival_object_uri"]}?resolve[]=digital_object'
        ).json()
        assert len(archival_object["instances"]) == 1
        tree = asnake_client.get(
            f'{archival_object["instances"][0]["digital_object"]["ref"]}/tree/root'
        ).json()
        for waypoint in tree["precomputed_waypoints"][""]["0"]:
            # ensure the s3 key matches the value in the file_uri_summary
            assert waypoint["file_uri_summary"].split(
                f's3://{config("PRESERVATION_BUCKET")}/'
            )[-1] in [s3_object["Key"] for s3_object in s3_response["Contents"]]
            # ensure the original filename from the s3 key matches the digital_object_component label
            assert waypoint["label"] in [
                s3_object["Key"].split("/")[-2] for s3_object in s3_response["Contents"]
            ]


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_s3_archivesspace_entries_multiple_multi_image_objects_3355ffxx1():
    return


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_s3_archivesspace_entries_multiple_multi_image_objects_3355ffxx2():
    return


def test_s3_archivesspace_entries_multiple_single_mixed_objects_5cae65(
    run, asnake_client, s3_client
):
    """Ensure single mixed format archival objects have entries in S3 and ArchivesSpace."""
    test_name = inspect.currentframe().f_code.co_name
    # NOTE "mixed" will always generate one each of audio, video, image
    # regardless of archival_object_count value
    run_output = run(
        test_name,
        collection_count=2,
        level="item",
        ancestors=[],
        timeout=120000,
    )
    # VALIDATE S3 UPLOADS AND ARCHIVESSPACE RECORDS
    for i in run_output:
        # get a list of s3 objects under this collection prefix
        s3_response = s3_client.list_objects_v2(
            Bucket=config("PRESERVATION_BUCKET"), Prefix=i["id_0"]
        )
        # get the digital_object tree
        archival_object = asnake_client.get(
            f'{i["archival_object_uri"]}?resolve[]=digital_object'
        ).json()
        assert len(archival_object["instances"]) == 1
        tree = asnake_client.get(
            f'{archival_object["instances"][0]["digital_object"]["ref"]}/tree/root'
        ).json()
        for waypoint in tree["precomputed_waypoints"][""]["0"]:
            # ensure the s3 key matches the value in the file_uri_summary
            assert waypoint["file_uri_summary"].split(
                f's3://{config("PRESERVATION_BUCKET")}/'
            )[-1] in [s3_object["Key"] for s3_object in s3_response["Contents"]]
            # ensure the original filename from the s3 key matches the digital_object_component label
            assert waypoint["label"] in [
                s3_object["Key"].split("/")[-2] for s3_object in s3_response["Contents"]
            ]


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_s3_archivesspace_entries_multiple_single_mixed_objects_5cae65xx1():
    return


@pytest.mark.skip(reason="placeholder for deletion of multiple collection records")
def test_s3_archivesspace_entries_multiple_single_mixed_objects_5cae65xx2():
    return


@pytest.mark.skipif(
    not os.getenv("RUN_TAPE_TESTS"),
    reason="environment variable RUN_TAPE_TESTS is not set",
)
def test_tape_reuse_top_container_records_d3bym(page: Page, asnake_client):
    """Items on the same tape should use the same top container record."""
    test_name = inspect.currentframe().f_code.co_name
    test_id = test_name.split("_")[-1]
    # MOVE TEST FILES TO INITIAL_ORIGINAL_FILES DIRECTORY
    move_test_files_to_initial_original_files_directory(
        "test_tape_reuse_top_container_records_d3by1"
    )
    move_test_files_to_initial_original_files_directory(
        "test_tape_reuse_top_container_records_d3by2"
    )
    # DELETE ANY EXISTING TEST RECORDS
    delete_archivesspace_test_records(asnake_client, test_id)
    # CREATE RESOURCE RECORD
    resource_create_response = create_archivesspace_test_resource(
        asnake_client, test_name, test_id
    )
    print(f"🐞 resource_create_response:{test_id}", resource_create_response.json())
    # CREATE ARCHIVAL OBJECT ITEM RECORD
    (
        item_create_response,
        item_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client,
        "test_tape_reuse_top_container_records_d3by1",
        "d3by1",
        resource_create_response.json()["uri"],
    )
    print("🐞 item2_create_response:d3by1", item_create_response.json())
    # CUSTOMIZE ARCHIVAL OBJECT ITEM RECORD
    item = asnake_client.get(item_create_response.json()["uri"]).json()
    item["title"] = "Item d3by1"
    item["component_id"] = "item-test_tape_reuse_top_container_records_d3by1"
    item_update_response = asnake_client.post(item["uri"], json=item)
    print("🐞 item_update_response:d3by1", item_update_response.json())
    # CREATE ARCHIVAL OBJECT ITEM2 RECORD
    (
        item2_create_response,
        item2_component_id,
    ) = create_archivesspace_test_archival_object_item(
        asnake_client,
        "test_tape_reuse_top_container_records_d3by2",
        "d3by2",
        resource_create_response.json()["uri"],
    )
    print("🐞 item2_create_response:d3by2", item2_create_response.json())
    # CUSTOMIZE ARCHIVAL OBJECT ITEM2 RECORD
    item2 = asnake_client.get(item2_create_response.json()["uri"]).json()
    item2["title"] = "Item d3by2"
    item2["component_id"] = "item-test_tape_reuse_top_container_records_d3by2"
    item2_update_response = asnake_client.post(item2["uri"], json=item2)
    print("🐞 item2_update_response:d3by2", item2_update_response.json())
    # RUN DISTILLERY TAPE PROCESS
    # NOTE increase timeout because tape drive can be quite slow to start up
    run_distillery(page, ["onsite"], timeout=300000)
    # VALIDATE TOP CONTAINER RECORDS
    # get the top_container uri of each item and compare them
    item = asnake_client.get(
        item_create_response.json()["uri"],
        params={"resolve[]": "digital_object", "resolve[]": "linked_agents"},
    ).json()
    for instance in item["instances"]:
        if instance.get("sub_container"):
            item_top_container_uri = instance["sub_container"]["top_container"]["ref"]
    item2 = asnake_client.get(
        item2_create_response.json()["uri"],
        params={"resolve[]": "digital_object", "resolve[]": "linked_agents"},
    ).json()
    for instance in item2["instances"]:
        if instance.get("sub_container"):
            item2_top_container_uri = instance["sub_container"]["top_container"]["ref"]
    # the linked top_container record should be the same for each item
    assert item_top_container_uri == item2_top_container_uri


def test_oralhistories_single_transcript_upload_publish_b1620f(
    page: Page, asnake_client, s3_client, timestamp
):
    """Upload a DOCX file and publish a transcript."""
    test_name = inspect.currentframe().f_code.co_name
    # BEGIN ORAL HISTORIES PROCESS
    kwargs = {
        "collection_count": 1,
        "archival_object_count": 1,
        "level": "file",
        "ancestors": [],
        "digital_file_count": 1,
    }
    # kwargs: collection_count
    delete_archivesspace_records(test_name, asnake_client, **kwargs)
    delete_filesystem_files()
    # kwargs: collection_count, archival_object_count, level, ancestors
    record_data = generate_archivesspace_records(test_name, asnake_client, **kwargs)
    print(f"🐞 record_data:{test_name}", record_data)
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    filesystem_files = generate_filesystem_files(test_name, **kwargs)
    # VALIDATE
    for _ in record_data:
        print(
            "🐞 ARCHIVAL OBJECT RECORD:",
            f'{config("ASPACE_STAFF_URL").rstrip("/")}/resolve/readonly?uri={_["archival_object_uri"]}',
        )
        # DELETE GITHUB TRANSCRIPTS
        # https://stackoverflow.com/a/72553300
        tmp_oralhistories = tempfile.mkdtemp()
        git_repo = git.Repo.clone_from(
            f'git@github.com:{config("ORALHISTORIES_GITHUB_REPO")}.git',
            tmp_oralhistories,
            env={"GIT_SSH_COMMAND": f'ssh -i {config("ORALHISTORIES_GITHUB_SSH_KEY")}'},
        )
        if os.path.exists(f'{tmp_oralhistories}/transcripts/{_["component_id"]}'):
            git_repo.index.remove(
                [f'transcripts/{_["component_id"]}'], working_tree=True, r=True
            )
            git_repo.index.commit(f'🤖 delete {_["component_id"]}')
            git_repo.remotes.origin.push()
        # DELETE S3 OBJECTS
        s3_response = s3_client.list_objects_v2(
            Bucket=config("ORALHISTORIES_BUCKET"), Prefix=_["component_id"]
        )
        print("🐞 s3_client.list_objects_v2", s3_response)
        if s3_response.get("Contents"):
            s3_keys = [
                {"Key": s3_object["Key"]} for s3_object in s3_response["Contents"]
            ]
            s3_response = s3_client.delete_objects(
                Bucket=config("ORALHISTORIES_BUCKET"), Delete={"Objects": s3_keys}
            )
        if config("RESOLVER_BUCKET", default=""):
            resolver_s3_response = s3_client.list_objects_v2(
                Bucket=config("RESOLVER_BUCKET"),
                Prefix="{}:{}".format(
                    config("RESOLVER_ORALHISTORIES_URL_PATH_PREFIX"),
                    _["component_id"],
                ),
            )
            print("🐞 resolver_s3_response", resolver_s3_response)
            if resolver_s3_response.get("Contents"):
                s3_keys = [
                    {"Key": s3_object["Key"]}
                    for s3_object in resolver_s3_response["Contents"]
                ]
                resolver_s3_response = s3_client.delete_objects(
                    Bucket=config("RESOLVER_BUCKET"), Delete={"Objects": s3_keys}
                )
        # RUN ORALHISTORIES PROCESSES
        # upload transcript
        print("🐞 filesystem_files", filesystem_files)
        run_oralhistories_add(page, filesystem_files[0])
        # wait for files to be updated by GitHub Actions
        assert wait_for_oralhistories_generated_files(git_repo)
        print(
            "🐞 GITHUB TRANSCRIPT:",
            f'https://github.com/{config("ORALHISTORIES_GITHUB_REPO")}/tree/main/transcripts/{_["component_id"]}',
        )
        # publish transcript
        run_oralhistories_publish(page, _["component_id"])
        # INVALIDATE CLOUDFRONT ITEMS
        if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
            invalidate_cloudfront_path(caller_reference=timestamp)
        # VALIDATE RESOLVER URL & WEB TRANSCRIPT
        if config("RESOLVER_BUCKET", default=""):
            resolver_url = "{}/{}:{}".format(
                config("RESOLVER_SERVICE_ENDPOINT").rstrip("/"),
                config("RESOLVER_ORALHISTORIES_URL_PATH_PREFIX"),
                _["component_id"],
            )
            print("🐞 RESOLVER URL:", resolver_url)
            page.goto(resolver_url)
        else:
            alchemist_url = "/".join(
                [
                    config("ALCHEMIST_BASE_URL").rstrip("/"),
                    config("ORALHISTORIES_URL_PATH_PREFIX"),
                    _["component_id"],
                ]
            )
            print("🐞 ALCHEMIST URL:", alchemist_url)
            page.goto(alchemist_url)
        # AWS S3 adds a trailing slash when it redirects to the index.html file.
        expect(page).to_have_url(
            "/".join(
                [
                    config("ALCHEMIST_BASE_URL").rstrip("/"),
                    config("ORALHISTORIES_URL_PATH_PREFIX"),
                    _["component_id"],
                    "",
                ]
            )
        )
        expect(page.get_by_text("access the full archival record")).to_have_attribute(
            "href",
            f'{config("ASPACE_PUBLIC_URL").rstrip("/")}{_["archival_object_uri"]}',
        )


def test_oralhistories_single_transcript_upload_update_publish_c6a18c(
    page: Page, asnake_client, s3_client, timestamp
):
    """Upload a DOCX file, update metadata, and publish transcript."""
    test_name = inspect.currentframe().f_code.co_name
    # BEGIN ORAL HISTORIES PROCESS
    kwargs = {
        "collection_count": 1,
        "archival_object_count": 1,
        "level": "file",
        "ancestors": [],
        "digital_file_count": 1,
    }
    # kwargs: collection_count
    delete_archivesspace_records(test_name, asnake_client, **kwargs)
    delete_filesystem_files()
    # kwargs: collection_count, archival_object_count, level, ancestors
    record_data = generate_archivesspace_records(test_name, asnake_client, **kwargs)
    print(f"🐞 record_data:{test_name}", record_data)
    # kwargs: collection_count, archival_object_count, level, digital_file_count
    filesystem_files = generate_filesystem_files(test_name, **kwargs)
    # VALIDATE
    for _ in record_data:
        print(
            "🐞 ARCHIVAL OBJECT RECORD:",
            f'{config("ASPACE_STAFF_URL").rstrip("/")}/resolve/readonly?uri={_["archival_object_uri"]}',
        )
        # DELETE GITHUB TRANSCRIPTS
        # https://stackoverflow.com/a/72553300
        tmp_oralhistories = tempfile.mkdtemp()
        git_repo = git.Repo.clone_from(
            f'git@github.com:{config("ORALHISTORIES_GITHUB_REPO")}.git',
            tmp_oralhistories,
            env={"GIT_SSH_COMMAND": f'ssh -i {config("ORALHISTORIES_GITHUB_SSH_KEY")}'},
        )
        if os.path.exists(f'{tmp_oralhistories}/transcripts/{_["component_id"]}'):
            git_repo.index.remove(
                [f'transcripts/{_["component_id"]}'], working_tree=True, r=True
            )
            git_repo.index.commit(f'🤖 delete {_["component_id"]}')
            git_repo.remotes.origin.push()
        # DELETE S3 OBJECTS
        s3_response = s3_client.list_objects_v2(
            Bucket=config("ORALHISTORIES_BUCKET"), Prefix=_["component_id"]
        )
        print("🐞 s3_client.list_objects_v2", s3_response)
        if s3_response.get("Contents"):
            s3_keys = [
                {"Key": s3_object["Key"]} for s3_object in s3_response["Contents"]
            ]
            s3_response = s3_client.delete_objects(
                Bucket=config("ORALHISTORIES_BUCKET"), Delete={"Objects": s3_keys}
            )
        if config("RESOLVER_BUCKET", default=""):
            resolver_s3_response = s3_client.list_objects_v2(
                Bucket=config("RESOLVER_BUCKET"),
                Prefix="{}:{}".format(
                    config("RESOLVER_ORALHISTORIES_URL_PATH_PREFIX"),
                    _["component_id"],
                ),
            )
            print("🐞 resolver_s3_response", resolver_s3_response)
            if resolver_s3_response.get("Contents"):
                s3_keys = [
                    {"Key": s3_object["Key"]}
                    for s3_object in resolver_s3_response["Contents"]
                ]
                resolver_s3_response = s3_client.delete_objects(
                    Bucket=config("RESOLVER_BUCKET"), Delete={"Objects": s3_keys}
                )
        # RUN ORALHISTORIES PROCESSES
        # upload transcript
        print("🐞 filesystem_files", filesystem_files)
        run_oralhistories_add(page, filesystem_files[0])
        # wait for files to be updated by GitHub Actions
        assert wait_for_oralhistories_generated_files(git_repo)
        print(
            "🐞 GITHUB TRANSCRIPT:",
            f'https://github.com/{config("ORALHISTORIES_GITHUB_REPO")}/tree/main/transcripts/{_["component_id"]}',
        )
        # update metadata
        generate = DocumentGenerator()
        abstract = generate.paragraph()
        item = asnake_client.get(_["archival_object_uri"]).json()
        item["notes"] = [
            {
                "jsonmodel_type": "note_singlepart",
                "type": "abstract",
                "content": [abstract],
                "publish": True,
            }
        ]
        item_update_response = asnake_client.post(item["uri"], json=item)
        print(f"🐞 item_update_response:{test_name}", item_update_response.json())
        run_oralhistories_update(page, _["component_id"])
        # wait for files to be updated by GitHub Actions
        assert wait_for_oralhistories_generated_files(
            git_repo, attempts=9, sleep_time=10
        )
        # publish transcript
        run_oralhistories_publish(page, _["component_id"])
        # INVALIDATE CLOUDFRONT ITEMS
        if config("ALCHEMIST_CLOUDFRONT_DISTRIBUTION_ID", default=False):
            invalidate_cloudfront_path(caller_reference=timestamp)
        # VALIDATE RESOLVER URL & WEB TRANSCRIPT
        if config("RESOLVER_BUCKET", default=""):
            resolver_url = "{}/{}:{}".format(
                config("RESOLVER_SERVICE_ENDPOINT").rstrip("/"),
                config("RESOLVER_ORALHISTORIES_URL_PATH_PREFIX"),
                _["component_id"],
            )
            print("🐞 RESOLVER URL:", resolver_url)
            page.goto(resolver_url)
        else:
            alchemist_url = "{}/{}/{}".format(
                config("ALCHEMIST_BASE_URL").rstrip("/"),
                config("ORALHISTORIES_URL_PATH_PREFIX"),
                _["component_id"],
            )
            print("🐞 ALCHEMIST URL:", alchemist_url)
            page.goto(alchemist_url)
        # AWS S3 adds a trailing slash when it redirects to the index.html file.
        expect(page).to_have_url(
            "{}/{}/{}/".format(
                config("ALCHEMIST_BASE_URL").rstrip("/"),
                config("ORALHISTORIES_URL_PATH_PREFIX"),
                _["component_id"],
            )
        )
        expect(page.get_by_text("access the full archival record")).to_have_attribute(
            "href",
            f'{config("ASPACE_PUBLIC_URL").rstrip("/")}{_["archival_object_uri"]}',
        )
        expect(page.locator("body")).to_contain_text(abstract[0:50])


# TODO
# def test_oralhistories_single_transcript_upload_edit_publish_5de292():
#     """Upload a DOCX file, edit markdown, and publish transcript."""
